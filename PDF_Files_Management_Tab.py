from PyQt5.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QListWidget, QStackedWidget, QPushButton, 
                             QLabel, QLineEdit, QComboBox, QSpinBox, QCheckBox, QFileDialog, QGroupBox, 
                             QMessageBox, QProgressBar, QSlider, QFontComboBox, QColorDialog, QGridLayout, QDialog)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QObject
import PyPDF2, pdfplumber, fitz
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import glob
from fontTools.ttLib import TTFont as TTFontCheck
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
import sys, os, io, subprocess, ctypes, re, unicodedata, cv2, itertools
from pytesseract import image_to_string
from PIL import Image
import numpy as np




class PDFManagementTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()
        self.check_ghostscript()

    def init_ui(self):
        layout = QHBoxLayout()
        self.operation_list = QListWidget()
        operations = [
                "ادغام پی دی اف ها", "تقسیم پی دی اف", "مرتب‌سازی صفحات" , "حذف صفحات", "استخراج صفحات",
                "قفل کردن پی دی اف", "باز کردن قفل پی دی اف", "چرخاندن صفحات", "فشرده‌سازی پی دی اف", 
                "استخراج متن", "استخراج تصاویر", "تصویر به پی دی اف", "اسکن PDFهای تصویری" 
            ]
        self.operation_list.addItems(operations)
        self.operation_list.currentRowChanged.connect(self.display_operation)
        layout.addWidget(self.operation_list, 1)

        self.stack = QStackedWidget()
        layout.addWidget(self.stack, 3)

        self.setLayout(layout)

        # Initialize widgets for each operation
        self.widgets = [
            self.create_merge_pdfs_widget(), self.create_split_pdf_widget(), self.create_reorder_pages_widget(),
            self.create_delete_pages_widget(), self.create_extract_pages_widget(),
            self.create_lock_pdf_widget(), self.create_unlock_pdf_widget(),
            self.create_rotate_pages_widget(), self.create_compress_pdf_widget(),
            self.create_extract_text_widget(), self.create_extract_images_widget(),
            self.create_image_to_pdf_widget(), self.create_scan_pdf_widget(),
        ]
        for widget in self.widgets:
            self.stack.addWidget(widget)

    def set_persian_font(self, widget):
        """تنظیم فونت فارسی برای ویجت و فرزندان آن"""
        font = widget.font()
        font.setFamily("B Nazanin")  # فونت فارسی
        font.setPointSize(10)
        widget.setFont(font)
        widget.setLayoutDirection(Qt.RightToLeft)  # تنظیم جهت RTL
        for child in widget.findChildren(QWidget):
            child.setFont(font)
            child.setLayoutDirection(Qt.RightToLeft)  # تنظیم جهت RTL برای فرزندان


    def parse_pages_input(self, input_str, total_pages):
        """
        ورودی کاربر را پارس می‌کند و لیستی از صفحات تکی یا بازه‌ها را برمی‌گرداند.
        پشتیبانی از ورودی‌هایی مثل: 1و2و3تا5 یا 1,2,4-6 یا 1تا3و1و2و4تا6
        """
        entries = []
        # جدا کردن بخش‌ها با استفاده از جداکننده‌های مختلف
        parts = re.split(r'[,\.\sو]+', input_str.strip())

        for part in parts:
            part = part.strip()
            if not part:
                continue
            # بررسی بازه‌ها (مثل 1تا3 یا 4-6)
            if 'تا' in part or '-' in part:
                range_parts = part.replace('تا', '-').split('-')
                if len(range_parts) == 2 and all(p.strip().isdigit() for p in range_parts):
                    start, end = map(int, range_parts)
                    # پشتیبانی از بازه‌های برعکس (مثل 5تا3)
                    start, end = min(start, end), max(start, end)
                    if 1 <= start <= total_pages and 1 <= end <= total_pages:
                        entries.append((start, end))
                    else:
                        raise ValueError(f"بازه {part} خارج از محدوده صفحات (1 تا {total_pages}) است.")
                else:
                    raise ValueError(f"فرمت بازه {part} نامعتبر است.")
            # بررسی صفحات تکی
            elif part.isdigit():
                page = int(part)
                if 1 <= page <= total_pages:
                    entries.append(page)
                else:
                    raise ValueError(f"صفحه {page} خارج از محدوده صفحات (1 تا {total_pages}) است.")
            else:
                raise ValueError(f"ورودی {part} نامعتبر است.")

        return entries


    def display_operation(self, index):
        self.stack.setCurrentIndex(index)



    def create_merge_pdfs_widget(self):
        # ایجاد ویجت اصلی
        widget = QWidget()
        widget.setLayoutDirection(Qt.RightToLeft)  # راست‌چین کردن ویجت
        main_layout = QVBoxLayout()
        main_layout.setAlignment(Qt.AlignCenter)
        main_layout.setSpacing(20)  # فاصله بهینه بین اجزا
        main_layout.setContentsMargins(30, 30, 30, 30)  # حاشیه‌های مدرن

        # استایل‌های مشترک
        input_style = """
            QLineEdit {
                padding: 10px;
                border-radius: 8px;
                border: 2px solid #e0e0e0;
                background-color: #ffffff;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QLineEdit:hover {
                border: 2px solid #007BFF;
                background-color: #f8faff;
            }
            QLineEdit:focus {
                border: 2px solid #0056b3;
                background-color: #ffffff;
            }
        """
        button_style = """
            QPushButton {
                padding: 10px;
                border-radius: 8px;
                background-color: #007BFF;
                color: white;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                font-weight: bold;
                border: none;
                min-width: 120px;
            }
            QPushButton:hover {
                background-color: #0056b3;
                border: 1px solid #004085;
            }
            QPushButton:pressed {
                background-color: #003d80;
            }
        """
        label_style = """
            QLabel {
                font-family: 'B Nazanin', 'Arial';
                font-size: 16px;
                font-weight: bold;
                color: #333333;
            }
        """
        group_style = """
            QGroupBox {
                border: 2px solid #e0e0e0;
                border-radius: 10px;
                padding: 15px;
                background-color: #ffffff;
                margin-top: 10px;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top Left;
                padding: 5px 10px;
                color: #333333;
                font-weight: bold;
            }
        """
        list_style = """
            QListWidget {
                padding: 10px;
                border-radius: 8px;
                border: 2px solid #e0e0e0;
                background-color: #ffffff;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QListWidget::item:hover {
                background-color: #f0f8ff;
            }
            QListWidget::item:selected {
                background-color: #007BFF;
                color: white;
            }
        """
        merge_button_style = """
            QPushButton {
                padding: 12px;
                border-radius: 10px;
                background-color: #28A745;
                color: white;
                font-family: 'B Nazanin', 'Arial';
                font-size: 16px;
                font-weight: bold;
                border: none;
                min-width: 200px;
            }
            QPushButton:hover {
                background-color: #218838;
                border: 1px solid #1e7e34;
            }
            QPushButton:pressed {
                background-color: #1a6b2d;
            }
        """

        # تنظیم فونت و استایل کلی
        widget.setStyleSheet("""
            QWidget {
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                background-color: #f5f5f5;
            }
        """)

        # گروه فایل‌های ورودی
        input_group = QGroupBox("فایل‌های PDF ورودی")
        input_group.setAlignment(Qt.AlignLeft)
        input_group.setStyleSheet(group_style)
        input_layout = QVBoxLayout()
        input_layout.setSpacing(10)

        input_label = QLabel("فایل‌های PDF ورودی:")
        input_label.setAlignment(Qt.AlignLeft)
        input_label.setStyleSheet(label_style)
        input_label.setToolTip("فایل‌های PDF را برای ادغام انتخاب کنید")

        self.merge_list = QListWidget()
        self.merge_list.setLayoutDirection(Qt.RightToLeft)
        self.merge_list.setStyleSheet(list_style)
        self.merge_list.setMinimumHeight(150)
        self.merge_list.setToolTip("لیست فایل‌های PDF انتخاب‌شده برای ادغام")

        button_layout = QHBoxLayout()
        button_layout.setSpacing(10)
        add_btn = QPushButton("افزودن فایل")
        add_btn.setStyleSheet(button_style)
        add_btn.setToolTip("افزودن فایل‌های PDF جدید")
        add_btn.clicked.connect(lambda: self.merge_list.addItems(QFileDialog.getOpenFileNames(self, "انتخاب فایل‌های PDF", "", "فایل‌های PDF (*.pdf)")[0]))

        remove_btn = QPushButton("حذف فایل")
        remove_btn.setStyleSheet(button_style)
        remove_btn.setToolTip("حذف فایل انتخاب‌شده از لیست")
        remove_btn.clicked.connect(lambda: [self.merge_list.takeItem(self.merge_list.row(item)) for item in self.merge_list.selectedItems()])

        up_btn = QPushButton("بالا")
        up_btn.setStyleSheet(button_style)
        up_btn.setToolTip("جابجایی فایل انتخاب‌شده به بالا")
        up_btn.clicked.connect(self.move_up)

        down_btn = QPushButton("پایین")
        down_btn.setStyleSheet(button_style)
        down_btn.setToolTip("جابجایی فایل انتخاب‌شده به پایین")
        down_btn.clicked.connect(self.move_down)

        button_layout.addWidget(add_btn)
        button_layout.addWidget(remove_btn)
        button_layout.addWidget(up_btn)
        button_layout.addWidget(down_btn)
        input_layout.addWidget(input_label)
        input_layout.addWidget(self.merge_list)
        input_layout.addLayout(button_layout)
        input_group.setLayout(input_layout)

        # گروه فایل خروجی
        output_group = QGroupBox("فایل PDF خروجی")
        output_group.setAlignment(Qt.AlignLeft)
        output_group.setStyleSheet(group_style)
        output_layout = QHBoxLayout()
        output_layout.setSpacing(10)

        output_label = QLabel("مسیر فایل خروجی:")
        output_label.setAlignment(Qt.AlignLeft)
        output_label.setStyleSheet(label_style)
        output_label.setToolTip("مسیر ذخیره فایل PDF ادغام‌شده")

        self.merge_output = QLineEdit()
        self.merge_output.setPlaceholderText("مسیر ذخیره فایل PDF خروجی (مثال: output.pdf)")
        self.merge_output.setAlignment(Qt.AlignLeft)
        self.merge_output.setStyleSheet(input_style)
        self.merge_output.setToolTip("مسیر فایل PDF ادغام‌شده را وارد یا انتخاب کنید")

        browse_output = QPushButton("انتخاب مسیر")
        browse_output.setStyleSheet(button_style)
        browse_output.setToolTip("انتخاب مسیر ذخیره فایل خروجی")
        browse_output.clicked.connect(lambda: self.merge_output.setText(QFileDialog.getSaveFileName(self, "ذخیره فایل PDF", "", "فایل‌های PDF (*.pdf)")[0]))

        output_layout.addWidget(output_label)
        output_layout.addWidget(self.merge_output)
        output_layout.addWidget(browse_output)
        output_group.setLayout(output_layout)

        # دکمه ادغام
        merge_btn = QPushButton("ادغام فایل‌های PDF")
        merge_btn.setStyleSheet(merge_button_style)
        merge_btn.setToolTip("شروع فرآیند ادغام فایل‌های PDF")
        merge_btn.clicked.connect(self.merge_pdfs)

        # افزودن اجزا به چیدمان اصلی
        main_layout.addWidget(input_group)
        main_layout.addWidget(output_group)
        main_layout.addWidget(merge_btn, alignment=Qt.AlignCenter)
        main_layout.addStretch()  # فضای خالی برای زیبایی

        widget.setLayout(main_layout)

        # تنظیم فونت فارسی (در صورت وجود تابع)
        if hasattr(self, 'set_persian_font'):
            self.set_persian_font(widget)

        return widget

    def move_up(self):
        current_row = self.merge_list.currentRow()
        if current_row > 0:
            item = self.merge_list.takeItem(current_row)
            self.merge_list.insertItem(current_row - 1, item)
            self.merge_list.setCurrentRow(current_row - 1)

    def move_down(self):
        current_row = self.merge_list.currentRow()
        if current_row < self.merge_list.count() - 1:
            item = self.merge_list.takeItem(current_row)
            self.merge_list.insertItem(current_row + 1, item)
            self.merge_list.setCurrentRow(current_row + 1)

    def merge_pdfs(self):
        output = self.merge_output.text().strip()
        pdfs = [self.merge_list.item(i).text() for i in range(self.merge_list.count())]

        # بررسی ورودی‌ها
        if not output or not pdfs:
            msg = QMessageBox()
            msg.setWindowTitle("هشدار")
            msg.setText("لطفاً مسیر فایل خروجی را مشخص کنید و حداقل یک فایل PDF انتخاب کنید.")
            msg.setIcon(QMessageBox.Warning)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
            return

        if not output.endswith('.pdf'):
            msg = QMessageBox()
            msg.setWindowTitle("هشدار")
            msg.setText("مسیر فایل خروجی باید با فرمت PDF باشد (مثال: output.pdf).")
            msg.setIcon(QMessageBox.Warning)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
            return

        try:
            merger = PyPDF2.PdfMerger()
            for pdf in pdfs:
                merger.append(pdf)
            with open(output, 'wb') as out:
                merger.write(out)
            merger.close()
            msg = QMessageBox()
            msg.setWindowTitle("موفقیت")
            msg.setText("فایل‌های PDF با موفقیت ادغام شدند.")
            msg.setIcon(QMessageBox.Information)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
        except FileNotFoundError:
            msg = QMessageBox()
            msg.setWindowTitle("خطا")
            msg.setText("یک یا چند فایل ورودی یافت نشدند.")
            msg.setIcon(QMessageBox.Critical)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
        except PermissionError:
            msg = QMessageBox()
            msg.setWindowTitle("خطا")
            msg.setText("دسترسی به فایل ممکن نیست. لطفاً مسیر دیگری انتخاب کنید.")
            msg.setIcon(QMessageBox.Critical)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
        except Exception as e:
            msg = QMessageBox()
            msg.setWindowTitle("خطا")
            msg.setText(f"خطا در ادغام فایل‌های PDF: {str(e)}")
            msg.setIcon(QMessageBox.Critical)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()



    def create_split_pdf_widget(self):
        # ایجاد ویجت اصلی
        widget = QWidget()
        widget.setLayoutDirection(Qt.RightToLeft)
        main_layout = QVBoxLayout()
        main_layout.setAlignment(Qt.AlignLeft)
        main_layout.setSpacing(20)
        main_layout.setContentsMargins(20, 20, 20, 20)

        # استایل‌های مشترک
        input_style = """
            QLineEdit {
                padding: 8px;
                border-radius: 5px;
                border: 1px solid #ccc;
                background-color: #fff;
            }
            QLineEdit:hover {
                border: 1px solid #007BFF;
                background-color: #f8faff;
            }
        """
        button_style = """
            QPushButton {
                padding: 8px;
                border-radius: 5px;
                background-color: #007BFF;
                color: white;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #0056b3;
                border: 1px solid #004085;
            }
            QPushButton:pressed {
                background-color: #003d80;
            }
        """
        split_button_style = """
            QPushButton {
                padding: 10px;
                border-radius: 5px;
                background-color: #28A745;
                color: white;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #218838;
                border: 1px solid #1e7e34;
            }
            QPushButton:pressed {
                background-color: #1a6b2d;
            }
        """
        label_style = "font-size: 14px; font-weight: bold; color: #333;"
        group_style = """
            QGroupBox {
                border: 1px solid #ddd;
                border-radius: 5px;
                padding: 10px;
                background-color: #f9f9f9;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top Left;
                padding: 5px 10px;
                color: #333333;
                font-weight: bold;
            }
        """
        combo_style = """
            QComboBox {
                padding: 8px;
                border-radius: 5px;
                border: 1px solid #ccc;
                background-color: #fff;
            }
            QComboBox:hover {
                border: 1px solid #007BFF;
                background-color: #f8faff;
            }
        """
        spinbox_style = """
            QSpinBox {
                padding: 8px;
                border-radius: 5px;
                border: 1px solid #ccc;
                background-color: #fff;
            }
            QSpinBox:hover {
                border: 1px solid #007BFF;
                background-color: #f8faff;
            }
        """

        # تنظیم فونت و استایل کلی
        widget.setStyleSheet("""
            QWidget {
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                background-color: #f8f9fa;
            }
        """)

        # گروه ورودی PDF
        input_group = QGroupBox("فایل PDF ورودی")
        input_group.setAlignment(Qt.AlignLeft)
        input_group.setStyleSheet(group_style)
        input_layout = QVBoxLayout()
        input_layout.setSpacing(5)

        input_label = QLabel("فایل PDF ورودی:")
        input_label.setStyleSheet(label_style)
        input_label.setAlignment(Qt.AlignLeft)
        input_label.setToolTip("مسیر فایل PDF ورودی را مشخص کنید")

        self.split_input = QLineEdit()
        self.split_input.setPlaceholderText("مسیر فایل PDF را انتخاب کنید")
        self.split_input.setAlignment(Qt.AlignLeft)
        self.split_input.setStyleSheet(input_style)
        self.split_input.setToolTip("فایل PDF را برای تقسیم انتخاب کنید")

        browse_input = QPushButton("انتخاب فایل")
        browse_input.setStyleSheet(button_style)
        browse_input.setToolTip("انتخاب فایل PDF")
        browse_input.clicked.connect(lambda: self.split_input.setText(QFileDialog.getOpenFileName(self, "انتخاب PDF", "", "فایل‌های PDF (*.pdf)")[0]))

        input_layout.addWidget(input_label)
        input_layout.addWidget(self.split_input)
        input_layout.addWidget(browse_input)
        input_group.setLayout(input_layout)

        # گروه نوع تقسیم
        type_group = QGroupBox("نوع تقسیم")
        type_group.setAlignment(Qt.AlignLeft)
        type_group.setStyleSheet(group_style)
        type_layout = QVBoxLayout()
        type_layout.setSpacing(5)

        type_label = QLabel("نوع تقسیم:")
        type_label.setStyleSheet(label_style)
        type_label.setAlignment(Qt.AlignLeft)
        type_label.setToolTip("روش تقسیم PDF را انتخاب کنید")

        self.split_type = QComboBox()
        self.split_type.addItems(["بر اساس شماره صفحات", "بر اساس تعداد فایل‌ها", "بر اساس تعداد صفحات در هر فایل"])
        self.split_type.setStyleSheet(combo_style)
        self.split_type.setToolTip("روش تقسیم PDF را انتخاب کنید")

        self.split_type.setItemData(0, "PDF را در صفحات مشخص شده تقسیم می‌کند (مثلاً 1,3,5-7)", Qt.ToolTipRole)
        self.split_type.setItemData(1, "PDF را به تعداد فایل‌های مشخص شده تقسیم می‌کند", Qt.ToolTipRole)
        self.split_type.setItemData(2, "PDF را به فایل‌هایی با تعداد صفحات مشخص شده تقسیم می‌کند", Qt.ToolTipRole)

        self.split_stack = QStackedWidget()

        # ویجت شماره صفحات
        page_widget = QWidget()
        page_widget.setLayoutDirection(Qt.RightToLeft)
        page_layout = QVBoxLayout()
        page_layout.setSpacing(5)

        pages_label = QLabel("شماره صفحات یا بازه‌ها (مثال: 1,3,5-7 یا 1و2و3تا5):")
        pages_label.setStyleSheet(label_style)
        pages_label.setAlignment(Qt.AlignLeft)
        pages_label.setToolTip("شماره صفحات یا بازه‌ها برای تقسیم (می‌توانید از , یا . یا فضای خالی یا و یا - و تا استفاده کنید)")

        self.split_pages = QLineEdit()
        self.split_pages.setPlaceholderText("شماره صفحات یا بازه‌ها (مثال: 1,3,5-7 یا 1و2و3تا5)")
        self.split_pages.setAlignment(Qt.AlignLeft)
        self.split_pages.setStyleSheet(input_style)
        self.split_pages.setToolTip("شماره صفحات یا بازه‌هایی که می‌خواهید PDF در آن‌ها تقسیم شود را وارد کنید")

        page_layout.addWidget(pages_label)
        page_layout.addWidget(self.split_pages)
        page_widget.setLayout(page_layout)

        # ویجت تعداد فایل‌ها
        files_widget = QWidget()
        files_widget.setLayoutDirection(Qt.RightToLeft)
        files_layout = QVBoxLayout()
        files_layout.setSpacing(5)

        num_files_label = QLabel("تعداد فایل‌ها:")
        num_files_label.setStyleSheet(label_style)
        num_files_label.setAlignment(Qt.AlignLeft)
        num_files_label.setToolTip("تعداد فایل‌های خروجی")

        self.split_num_files = QSpinBox()
        self.split_num_files.setMinimum(1)
        self.split_num_files.setStyleSheet(spinbox_style)
        self.split_num_files.setToolTip("تعداد فایل‌های خروجی مورد نظر را وارد کنید")

        files_layout.addWidget(num_files_label)
        files_layout.addWidget(self.split_num_files)
        files_widget.setLayout(files_layout)

        # ویجت تعداد صفحات در هر فایل
        pages_per_file_widget = QWidget()
        pages_per_file_widget.setLayoutDirection(Qt.RightToLeft)
        pages_per_file_layout = QVBoxLayout()
        pages_per_file_layout.setSpacing(5)

        pages_per_file_label = QLabel("تعداد صفحات در هر فایل:")
        pages_per_file_label.setStyleSheet(label_style)
        pages_per_file_label.setAlignment(Qt.AlignLeft)
        pages_per_file_label.setToolTip("تعداد صفحات در هر فایل")

        self.split_pages_per_file = QSpinBox()
        self.split_pages_per_file.setMinimum(1)
        self.split_pages_per_file.setStyleSheet(spinbox_style)
        self.split_pages_per_file.setToolTip("تعداد صفحات در هر فایل خروجی را وارد کنید")

        pages_per_file_layout.addWidget(pages_per_file_label)
        pages_per_file_layout.addWidget(self.split_pages_per_file)
        pages_per_file_widget.setLayout(pages_per_file_layout)

        self.split_stack.addWidget(page_widget)
        self.split_stack.addWidget(files_widget)
        self.split_stack.addWidget(pages_per_file_widget)
        self.split_type.currentIndexChanged.connect(self.split_stack.setCurrentIndex)

        type_layout.addWidget(type_label)
        type_layout.addWidget(self.split_type)
        type_layout.addWidget(self.split_stack)
        type_group.setLayout(type_layout)

        # گروه خروجی
        output_group = QGroupBox("خروجی")
        output_group.setAlignment(Qt.AlignLeft)
        output_group.setStyleSheet(group_style)
        output_layout = QVBoxLayout()
        output_layout.setSpacing(5)

        output_dir_label = QLabel("پوشه خروجی:")
        output_dir_label.setStyleSheet(label_style)
        output_dir_label.setAlignment(Qt.AlignLeft)
        output_dir_label.setToolTip("مسیر پوشه خروجی")

        self.split_output_dir = QLineEdit()
        self.split_output_dir.setPlaceholderText("مسیر پوشه خروجی را انتخاب کنید")
        self.split_output_dir.setAlignment(Qt.AlignLeft)
        self.split_output_dir.setStyleSheet(input_style)
        self.split_output_dir.setToolTip("پوشه‌ای که فایل‌های خروجی در آن ذخیره می‌شوند")

        browse_dir = QPushButton("انتخاب پوشه")
        browse_dir.setStyleSheet(button_style)
        browse_dir.setToolTip("انتخاب پوشه خروجی")
        browse_dir.clicked.connect(lambda: self.split_output_dir.setText(QFileDialog.getExistingDirectory(self, "انتخاب پوشه خروجی")))

        base_name_label = QLabel("نام پایه فایل‌ها:")
        base_name_label.setStyleSheet(label_style)
        base_name_label.setAlignment(Qt.AlignLeft)
        base_name_label.setToolTip("نام پایه فایل‌های خروجی")

        self.split_base_name = QLineEdit()
        self.split_base_name.setPlaceholderText("نام پایه فایل‌ها (مثال: document)")
        self.split_base_name.setAlignment(Qt.AlignLeft)
        self.split_base_name.setStyleSheet(input_style)
        self.split_base_name.setToolTip("نام پایه برای فایل‌های خروجی (مثال: document_بخش1.pdf)")

        output_layout.addWidget(output_dir_label)
        output_layout.addWidget(self.split_output_dir)
        output_layout.addWidget(browse_dir)
        output_layout.addWidget(base_name_label)
        output_layout.addWidget(self.split_base_name)
        output_group.setLayout(output_layout)

        # دکمه تقسیم
        split_btn = QPushButton("تقسیم PDF")
        split_btn.setStyleSheet(split_button_style)
        split_btn.setToolTip("شروع فرآیند تقسیم PDF")
        split_btn.clicked.connect(self.split_pdf)

        # افزودن گروه‌ها به چیدمان اصلی
        main_layout.addWidget(input_group)
        main_layout.addWidget(type_group)
        main_layout.addWidget(output_group)
        main_layout.addWidget(split_btn, alignment=Qt.AlignCenter)

        widget.setLayout(main_layout)
        return widget


    def split_pdf(self):
        input_file = self.split_input.text().strip()
        output_dir = self.split_output_dir.text().strip()
        base_name = self.split_base_name.text().strip()

        if not all([input_file, output_dir, base_name]):
            QMessageBox.warning(self, "هشدار", "لطفاً همه فیلدها را پر کنید.")
            return

        if not input_file.endswith('.pdf'):
            QMessageBox.warning(self, "هشدار", "لطفاً یک فایل PDF معتبر انتخاب کنید.")
            return

        try:
            with open(input_file, 'rb') as f:
                # استفاده از strict=False برای مدیریت فایل‌های PDF مشکل‌دار
                reader = PyPDF2.PdfReader(f, strict=False)
                total_pages = len(reader.pages)
                split_type = self.split_type.currentText()

                if split_type == "بر اساس شماره صفحات":
                    if not self.split_pages.text().strip():
                        QMessageBox.warning(self, "هشدار", "لطفاً شماره صفحات یا بازه‌ها را وارد کنید.")
                        return
                    try:
                        # پارس کردن ورودی کاربر
                        page_entries = self.parse_pages_input(self.split_pages.text(), total_pages)
                        
                        # ایجاد فایل‌های PDF برای هر ورودی (صفحه تکی یا بازه)
                        for i, entry in enumerate(page_entries, start=1):
                            writer = PyPDF2.PdfWriter()
                            if isinstance(entry, int):  # صفحه تکی
                                writer.add_page(reader.pages[entry - 1])
                                output_file = os.path.join(output_dir, f"{base_name}_صفحه{entry}.pdf")
                            else:  # بازه
                                start, end = entry
                                for page in range(start - 1, end):
                                    writer.add_page(reader.pages[page])
                                output_file = os.path.join(output_dir, f"{base_name}_صفحات{start}تا{end}.pdf")
                            
                            with open(output_file, 'wb') as out:
                                writer.write(out)

                    except ValueError as e:
                        QMessageBox.warning(self, "هشدار", str(e))
                        return

                elif split_type == "بر اساس تعداد فایل‌ها":
                    num_files = self.split_num_files.value()
                    if num_files < 1:
                        QMessageBox.warning(self, "هشدار", "تعداد فایل‌ها باید حداقل 1 باشد.")
                        return
                    if num_files > total_pages:
                        num_files = total_pages
                    pages_per_file = total_pages // num_files
                    extra = total_pages % num_files
                    start = 0
                    for i in range(num_files):
                        end = start + pages_per_file + (1 if i < extra else 0)
                        if start < end:
                            writer = PyPDF2.PdfWriter()
                            for page in range(start, end):
                                writer.add_page(reader.pages[page])
                            output_file = os.path.join(output_dir, f"{base_name}_بخش{i+1}.pdf")
                            with open(output_file, 'wb') as out:
                                writer.write(out)
                        start = end

                else:  # بر اساس تعداد صفحات در هر فایل
                    pages_per_file = self.split_pages_per_file.value()
                    if pages_per_file < 1:
                        QMessageBox.warning(self, "هشدار", "تعداد صفحات باید حداقل 1 باشد.")
                        return
                    for start in range(0, total_pages, pages_per_file):
                        end = min(start + pages_per_file, total_pages)
                        writer = PyPDF2.PdfWriter()
                        for page in range(start, end):
                            writer.add_page(reader.pages[page])
                        part = start // pages_per_file + 1
                        output_file = os.path.join(output_dir, f"{base_name}_بخش{part}.pdf")
                        with open(output_file, 'wb') as out:
                            writer.write(out)

            QMessageBox.information(self, "موفقیت", "PDF با موفقیت تقسیم شد.")

        except FileNotFoundError:
            QMessageBox.critical(self, "خطا", "فایل ورودی یافت نشد.")
        except PyPDF2.errors.PdfReadError:
            QMessageBox.critical(self, "خطا", "فایل PDF نامعتبر یا خراب است. لطفاً فایل دیگری انتخاب کنید.")
        except PermissionError:
            QMessageBox.critical(self, "خطا", "دسترسی به پوشه یا فایل ممکن نیست.")
        except Exception as e:
            QMessageBox.critical(self, "خطا", f"تقسیم PDF ناموفق بود: {str(e)}")



    def create_delete_pages_widget(self):
        # ایجاد ویجت اصلی
        widget = QWidget()
        widget.setLayoutDirection(Qt.RightToLeft)  # راست‌چین کردن کل ویجت
        main_layout = QVBoxLayout()
        main_layout.setAlignment(Qt.AlignLeft)
        main_layout.setSpacing(20)  # فاصله مناسب بین گروه‌ها
        main_layout.setContentsMargins(30, 30, 30, 30)  # حاشیه‌های مدرن

        # استایل‌های مشترک
        input_style = """
            QLineEdit {
                padding: 8px;
                border-radius: 5px;
                border: 1px solid #ccc;
                background-color: #fff;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QLineEdit:hover {
                border: 1px solid #007BFF;
                background-color: #f0f8ff;
            }
            QLineEdit:focus {
                border: 1px solid #007BFF;
                background-color: #f0f8ff;
            }
        """
        button_style = """
            QPushButton {
                padding: 8px;
                border-radius: 5px;
                background-color: #007BFF;
                color: white;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #0056b3;
            }
            QPushButton:pressed {
                background-color: #003d80;
            }
        """
        label_style = """
            font-family: 'B Nazanin', 'Arial';
            font-size: 14px;
            font-weight: bold;
            color: #333;
        """
        group_style = """
            QGroupBox {
                border: 1px solid #ddd;
                border-radius: 5px;
                padding: 10px;
                background-color: #fff;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top Left;
                padding: 0 5px;
                color: #333;
                font-weight: bold;
            }
        """
        delete_button_style = """
            QPushButton {
                padding: 10px;
                border-radius: 5px;
                background-color: #28A745;
                color: white;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #218838;
            }
            QPushButton:pressed {
                background-color: #1e7e34;
            }
        """

        # تنظیم فونت فارسی
        widget.setStyleSheet("""
            QWidget {
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                background-color: #f8f9fa;
            }
        """)

        # گروه ورودی فایل PDF
        input_group = QGroupBox("فایل PDF ورودی")
        input_group.setAlignment(Qt.AlignLeft)
        input_group.setStyleSheet(group_style)
        input_layout = QVBoxLayout()
        input_layout.setSpacing(10)

        input_label = QLabel("فایل PDF ورودی:")
        input_label.setAlignment(Qt.AlignLeft)
        input_label.setStyleSheet(label_style)
        input_label.setToolTip("فایل PDF را که می‌خواهید صفحات آن را حذف کنید، انتخاب کنید.")

        self.delete_input = QLineEdit()
        self.delete_input.setPlaceholderText("مسیر فایل PDF را انتخاب کنید")
        self.delete_input.setAlignment(Qt.AlignLeft)
        self.delete_input.setStyleSheet(input_style)
        self.delete_input.setToolTip("مسیر فایل PDF ورودی را وارد کنید یا از دکمه انتخاب فایل استفاده کنید.")

        browse_btn = QPushButton("انتخاب فایل")
        browse_btn.setStyleSheet(button_style)
        browse_btn.clicked.connect(self.browse_input_file_delete_pages)
        browse_btn.setToolTip("برای انتخاب فایل PDF، روی این دکمه کلیک کنید.")

        input_layout.addWidget(input_label)
        input_layout.addWidget(self.delete_input)
        input_layout.addWidget(browse_btn)
        input_group.setLayout(input_layout)

        # گروه صفحات برای حذف
        pages_group = QGroupBox("صفحات برای حذف")
        pages_group.setAlignment(Qt.AlignLeft)
        pages_group.setStyleSheet(group_style)
        pages_layout = QVBoxLayout()
        pages_layout.setSpacing(10)

        pages_label = QLabel("صفحات برای حذف (مثال: 1و2و3تا5 یا 4,8,6-24):")
        pages_label.setAlignment(Qt.AlignLeft)
        pages_label.setStyleSheet(label_style)
        pages_label.setToolTip("شماره صفحاتی که می‌خواهید حذف شوند را وارد کنید.")

        self.delete_pages = QLineEdit()
        self.delete_pages.setPlaceholderText("صفحات مورد نظر (مثال: 1و2و3تا5 یا 4,8,6-24)")
        self.delete_pages.setAlignment(Qt.AlignLeft)
        self.delete_pages.setStyleSheet(input_style)
        self.delete_pages.setToolTip("شماره صفحات را به صورت جدا شده با واو، کاما یا محدوده (با تا یا -) وارد کنید.")

        pages_layout.addWidget(pages_label)
        pages_layout.addWidget(self.delete_pages)
        pages_group.setLayout(pages_layout)

        # گروه خروجی فایل
        output_group = QGroupBox("فایل خروجی")
        output_group.setAlignment(Qt.AlignLeft)
        output_group.setStyleSheet(group_style)
        output_layout = QVBoxLayout()
        output_layout.setSpacing(10)

        output_label = QLabel("فایل خروجی:")
        output_label.setAlignment(Qt.AlignLeft)
        output_label.setStyleSheet(label_style)
        output_label.setToolTip("مسیر و نام فایل PDF خروجی را مشخص کنید.")

        self.delete_output = QLineEdit()
        self.delete_output.setPlaceholderText("مسیر ذخیره فایل خروجی")
        self.delete_output.setAlignment(Qt.AlignLeft)
        self.delete_output.setStyleSheet(input_style)
        self.delete_output.setToolTip("مسیر و نام فایل PDF خروجی را وارد کنید یا از دکمه انتخاب مسیر استفاده کنید.")

        browse_output = QPushButton("انتخاب مسیر ذخیره")
        browse_output.setStyleSheet(button_style)
        browse_output.clicked.connect(self.browse_output_file_delete_pages)
        browse_output.setToolTip("برای انتخاب مسیر ذخیره فایل خروجی، روی این دکمه کلیک کنید.")

        output_layout.addWidget(output_label)
        output_layout.addWidget(self.delete_output)
        output_layout.addWidget(browse_output)
        output_group.setLayout(output_layout)

        # دکمه حذف صفحات
        delete_btn = QPushButton("حذف صفحات")
        delete_btn.setStyleSheet(delete_button_style)
        delete_btn.clicked.connect(self.delete_pages_func)
        delete_btn.setToolTip("برای حذف صفحات و ذخیره فایل جدید، روی این دکمه کلیک کنید.")

        # افزودن گروه‌ها و دکمه به چیدمان اصلی
        main_layout.addWidget(input_group)
        main_layout.addWidget(pages_group)
        main_layout.addWidget(output_group)
        main_layout.addWidget(delete_btn, alignment=Qt.AlignCenter)

        widget.setLayout(main_layout)
        return widget

    def browse_input_file_delete_pages(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "انتخاب فایل PDF", "", "فایل‌های PDF (*.pdf)")
        if file_path:
            self.delete_input.setText(file_path)

    def browse_output_file_delete_pages(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "ذخیره فایل خروجی", "", "فایل‌های PDF (*.pdf)")
        if file_path:
            self.delete_output.setText(file_path)
            # بررسی فرمت فایل خروجی
            if not file_path.endswith('.pdf'):
                msg = QMessageBox()
                msg.setWindowTitle("هشدار")
                msg.setText("لطفاً مسیر فایل خروجی را با فرمت PDF انتخاب کنید.")
                msg.setIcon(QMessageBox.Warning)
                msg.setLayoutDirection(Qt.RightToLeft)
                msg.exec_()
                self.delete_output.setText("")  # پاک کردن مسیر نامعتبر
                return

    def delete_pages_func(self):
        input_file = self.delete_input.text().strip()
        output_file = self.delete_output.text().strip()
        pages_text = self.delete_pages.text().strip()

        # بررسی ورودی‌ها
        if not input_file or not output_file or not pages_text:
            msg = QMessageBox()
            msg.setWindowTitle("هشدار")
            msg.setText("لطفاً تمام فیلدها (ورودی، خروجی و صفحات) را پر کنید.")
            msg.setIcon(QMessageBox.Warning)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
            return

        if not input_file.endswith('.pdf') or not output_file.endswith('.pdf'):
            msg = QMessageBox()
            msg.setWindowTitle("هشدار")
            msg.setText("لطفاً فایل‌های PDF معتبر انتخاب کنید.")
            msg.setIcon(QMessageBox.Warning)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
            return

        try:
            # خواندن فایل PDF
            with open(input_file, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                total_pages = len(reader.pages)

                # پارس کردن ورودی صفحات
                entries = self.parse_pages_input(pages_text, total_pages)

                # جمع‌آوری صفحاتی که باید حذف شوند
                pages_to_delete = set()
                for entry in entries:
                    if isinstance(entry, int):
                        pages_to_delete.add(entry)
                    else:
                        start, end = entry
                        pages_to_delete.update(range(start, end + 1))

                if not pages_to_delete:
                    msg = QMessageBox()
                    msg.setWindowTitle("هشدار")
                    msg.setText("هیچ صفحه معتبری برای حذف مشخص نشده است.")
                    msg.setIcon(QMessageBox.Warning)
                    msg.setLayoutDirection(Qt.RightToLeft)
                    msg.exec_()
                    return

                # ایجاد فایل جدید بدون صفحات مشخص‌شده
                writer = PyPDF2.PdfWriter()
                for i in range(total_pages):
                    if i + 1 not in pages_to_delete:
                        writer.add_page(reader.pages[i])

                # ذخیره فایل خروجی
                with open(output_file, 'wb') as out:
                    writer.write(out)

                msg = QMessageBox()
                msg.setWindowTitle("موفقیت")
                msg.setText("صفحات با موفقیت حذف شدند.")
                msg.setIcon(QMessageBox.Information)
                msg.setLayoutDirection(Qt.RightToLeft)
                msg.exec_()

        except ValueError as ve:
            msg = QMessageBox()
            msg.setWindowTitle("هشدار")
            msg.setText(str(ve))
            msg.setIcon(QMessageBox.Warning)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
        except FileNotFoundError:
            msg = QMessageBox()
            msg.setWindowTitle("خطا")
            msg.setText("فایل ورودی یافت نشد.")
            msg.setIcon(QMessageBox.Critical)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
        except PermissionError:
            msg = QMessageBox()
            msg.setWindowTitle("خطا")
            msg.setText("دسترسی به فایل ممکن نیست. لطفاً مسیر دیگری انتخاب کنید.")
            msg.setIcon(QMessageBox.Critical)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
        except Exception as e:
            msg = QMessageBox()
            msg.setWindowTitle("خطا")
            msg.setText(f"خطا در حذف صفحات: {str(e)}")
            msg.setIcon(QMessageBox.Critical)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()




    def create_extract_pages_widget(self):
        # ایجاد ویجت اصلی
        widget = QWidget()
        widget.setLayoutDirection(Qt.RightToLeft)  # تنظیم جهت به راست‌به‌چپ
        main_layout = QVBoxLayout()
        main_layout.setAlignment(Qt.AlignCenter)  # تراز وسط برای جذابیت بیشتر
        main_layout.setSpacing(20)  # فاصله مناسب بین گروه‌ها
        main_layout.setContentsMargins(30, 30, 30, 30)  # حاشیه‌های مدرن و بهینه

        # استایل‌های مشترک
        input_style = """
            QLineEdit {
                padding: 10px;
                border-radius: 8px;
                border: 2px solid #e0e0e0;
                background-color: #ffffff;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QLineEdit:hover {
                border: 2px solid #007BFF;
                background-color: #f8faff;
            }
            QLineEdit:focus {
                border: 2px solid #0056b3;
                background-color: #ffffff;
            }
        """
        button_style = """
            QPushButton {
                padding: 10px;
                border-radius: 8px;
                background-color: #007BFF;
                color: white;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                font-weight: bold;
                border: none;
                min-width: 120px;
            }
            QPushButton:hover {
                background-color: #0056b3;
                border: 1px solid #004085;
            }
            QPushButton:pressed {
                background-color: #003d80;
            }
        """
        label_style = """
            QLabel {
                font-family: 'B Nazanin', 'Arial';
                font-size: 16px;
                font-weight: bold;
                color: #333333;
            }
        """
        group_style = """
            QGroupBox {
                border: 2px solid #e0e0e0;
                border-radius: 10px;
                padding: 15px;
                background-color: #ffffff;
                margin-top: 10px;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top Left;
                padding: 5px 10px;
                color: #333333;
                font-weight: bold;
            }
        """
        extract_button_style = """
            QPushButton {
                padding: 12px;
                border-radius: 10px;
                background-color: #28A745;
                color: white;
                font-family: 'B Nazanin', 'Arial';
                font-size: 16px;
                font-weight: bold;
                border: none;
                min-width: 200px;
            }
            QPushButton:hover {
                background-color: #218838;
                border: 1px solid #1e7e34;
            }
            QPushButton:pressed {
                background-color: #1a6b2d;
            }
        """

        # تنظیم فونت و استایل کلی
        widget.setStyleSheet("""
            QWidget {
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                background-color: #f5f5f5;
            }
        """)

        # گروه ورودی فایل PDF
        input_group = QGroupBox("فایل PDF ورودی")
        input_group.setAlignment(Qt.AlignLeft)
        input_group.setStyleSheet(group_style)
        input_layout = QVBoxLayout()
        input_layout.setSpacing(10)

        input_label = QLabel("فایل PDF ورودی:")
        input_label.setAlignment(Qt.AlignLeft)
        input_label.setStyleSheet(label_style)
        input_label.setToolTip("فایل PDF را که می‌خواهید صفحات آن را استخراج کنید، انتخاب کنید.")

        self.extract_input = QLineEdit()
        self.extract_input.setPlaceholderText("مسیر فایل PDF را انتخاب کنید")
        self.extract_input.setAlignment(Qt.AlignLeft)
        self.extract_input.setStyleSheet(input_style)
        self.extract_input.setToolTip("مسیر فایل PDF ورودی را وارد کنید یا از دکمه انتخاب فایل استفاده کنید.")

        browse_btn = QPushButton("انتخاب فایل")
        browse_btn.setStyleSheet(button_style)
        browse_btn.clicked.connect(lambda: self.extract_input.setText(QFileDialog.getOpenFileName(self, "انتخاب PDF", "", "فایل‌های PDF (*.pdf)")[0]))
        browse_btn.setToolTip("برای انتخاب فایل PDF، روی این دکمه کلیک کنید.")

        input_layout.addWidget(input_label)
        input_layout.addWidget(self.extract_input)
        input_layout.addWidget(browse_btn)
        input_group.setLayout(input_layout)

        # گروه صفحات برای استخراج
        pages_group = QGroupBox("صفحات برای استخراج")
        pages_group.setAlignment(Qt.AlignLeft)
        pages_group.setStyleSheet(group_style)
        pages_layout = QVBoxLayout()
        pages_layout.setSpacing(10)

        pages_label = QLabel("صفحات برای استخراج (مثال: 1و2و3تا5 یا 4,8,6-24):")
        pages_label.setAlignment(Qt.AlignLeft)
        pages_label.setStyleSheet(label_style)
        pages_label.setToolTip("شماره صفحاتی که می‌خواهید استخراج شوند را وارد کنید.")

        self.extract_pages = QLineEdit()
        self.extract_pages.setPlaceholderText("صفحات مورد نظر (مثال: 1و2و3تا5 یا 4,8,6-24)")
        self.extract_pages.setAlignment(Qt.AlignLeft)
        self.extract_pages.setStyleSheet(input_style)
        self.extract_pages.setToolTip("شماره صفحات را به صورت جدا شده با واو، کاما یا محدوده (با تا یا -) وارد کنید.")

        pages_layout.addWidget(pages_label)
        pages_layout.addWidget(self.extract_pages)
        pages_group.setLayout(pages_layout)

        # گروه خروجی فایل
        output_group = QGroupBox("فایل خروجی")
        output_group.setAlignment(Qt.AlignLeft)
        output_group.setStyleSheet(group_style)
        output_layout = QVBoxLayout()
        output_layout.setSpacing(10)

        output_label = QLabel("فایل خروجی:")
        output_label.setAlignment(Qt.AlignLeft)
        output_label.setStyleSheet(label_style)
        output_label.setToolTip("مسیر و نام فایل PDF خروجی را مشخص کنید.")

        self.extract_output = QLineEdit()
        self.extract_output.setPlaceholderText("مسیر ذخیره فایل خروجی")
        self.extract_output.setAlignment(Qt.AlignLeft)
        self.extract_output.setStyleSheet(input_style)
        self.extract_output.setToolTip("مسیر و نام فایل PDF خروجی را وارد کنید یا از دکمه انتخاب مسیر استفاده کنید.")

        browse_output = QPushButton("انتخاب مسیر ذخیره")
        browse_output.setStyleSheet(button_style)
        browse_output.clicked.connect(lambda: self.extract_output.setText(QFileDialog.getSaveFileName(self, "ذخیره فایل خروجی", "", "فایل‌های PDF (*.pdf)")[0]))
        browse_output.setToolTip("برای انتخاب مسیر ذخیره فایل خروجی، روی این دکمه کلیک کنید.")

        output_layout.addWidget(output_label)
        output_layout.addWidget(self.extract_output)
        output_layout.addWidget(browse_output)
        output_group.setLayout(output_layout)

        # دکمه استخراج
        extract_btn = QPushButton("استخراج صفحات")
        extract_btn.setStyleSheet(extract_button_style)
        extract_btn.clicked.connect(self.extract_pages_func)
        extract_btn.setToolTip("برای استخراج صفحات و ذخیره فایل جدید، روی این دکمه کلیک کنید.")

        # افزودن گروه‌ها و دکمه به چیدمان اصلی
        main_layout.addWidget(input_group)
        main_layout.addWidget(pages_group)
        main_layout.addWidget(output_group)
        main_layout.addWidget(extract_btn, alignment=Qt.AlignCenter)

        widget.setLayout(main_layout)
        return widget

    def extract_pages_func(self):
        input_file = self.extract_input.text().strip()
        output_file = self.extract_output.text().strip()
        pages_text = self.extract_pages.text().strip()

        # بررسی ورودی‌ها
        if not input_file or not output_file:
            QMessageBox.warning(self, "هشدار", "لطفاً فیلدهای ورودی و خروجی را پر کنید.")
            return

        if not input_file.endswith('.pdf') or not output_file.endswith('.pdf'):
            QMessageBox.warning(self, "هشدار", "لطفاً فایل‌های PDF معتبر انتخاب کنید.")
            return

        try:
            # خواندن فایل PDF
            with open(input_file, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                total_pages = len(reader.pages)

                # پارس کردن ورودی صفحات
                entries = self.parse_pages_input(pages_text, total_pages)

                # جمع‌آوری صفحاتی که باید استخراج شوند
                pages_to_extract = set()
                for entry in entries:
                    if isinstance(entry, int):
                        pages_to_extract.add(entry)
                    else:
                        start, end = entry
                        pages_to_extract.update(range(start, end + 1))

                if not pages_to_extract:
                    QMessageBox.warning(self, "هشدار", "هیچ صفحه معتبری برای استخراج مشخص نشده است.")
                    return

                # ایجاد فایل جدید با صفحات مشخص‌شده
                writer = PyPDF2.PdfWriter()
                for page_num in sorted(pages_to_extract):
                    writer.add_page(reader.pages[page_num - 1])

                # ذخیره فایل خروجی
                with open(output_file, 'wb') as out:
                    writer.write(out)

            QMessageBox.information(self, "موفقیت", "صفحات با موفقیت استخراج شدند.")
        except ValueError as ve:
            QMessageBox.warning(self, "هشدار", str(ve))
        except FileNotFoundError:
            QMessageBox.critical(self, "خطا", "فایل ورودی یافت نشد.")
        except PermissionError:
            QMessageBox.critical(self, "خطا", "دسترسی به فایل ممکن نیست. لطفاً مسیر دیگری انتخاب کنید.")
        except Exception as e:
            QMessageBox.critical(self, "خطا", f"خطا در استخراج صفحات: {str(e)}")




    def create_lock_pdf_widget(self):
        # ایجاد ویجت اصلی
        widget = QWidget()
        widget.setLayoutDirection(Qt.RightToLeft)  # تنظیم جهت راست‌به‌چپ
        main_layout = QVBoxLayout()
        main_layout.setAlignment(Qt.AlignLeft)  # ترازبندی به راست
        main_layout.setSpacing(20)  # فاصله مناسب بین گروه‌ها
        main_layout.setContentsMargins(30, 30, 30, 30)  # حاشیه‌های مدرن و بهینه

        # **استایل‌های مشترک**
        input_style = """
            QLineEdit {
                padding: 8px;
                border-radius: 5px;
                border: 1px solid #ccc;
                background-color: #fff;
            }
            QLineEdit:hover {
                border: 1px solid #007BFF;
                background-color: #f0f8ff;
            }
            QLineEdit:focus {
                border: 1px solid #007BFF;
                background-color: #f0f8ff;
            }
        """
        button_style = """
            QPushButton {
                padding: 8px;
                border-radius: 5px;
                background-color: #007BFF;
                color: white;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #0056b3;
            }
            QPushButton:pressed {
                background-color: #003d80;
            }
        """
        label_style = "font-size: 14px; font-weight: bold; color: #333;"
        checkbox_style = "font-size: 14px; color: #333; padding: 5px;"
        group_style = """
            QGroupBox {
                border: 1px solid #ddd;
                border-radius: 5px;
                padding: 10px;
                background-color: #fff;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top Left;
                padding: 0 5px;
                color: #333;
                font-weight: bold;
            }
        """
        lock_button_style = """
            QPushButton {
                padding: 10px;
                border-radius: 5px;
                background-color: #28A745;
                color: white;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #218838;
            }
            QPushButton:pressed {
                background-color: #1e7e34;
            }
        """

        # **تنظیم فونت فارسی و استایل کلی**
        widget.setStyleSheet("""
            QWidget {
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                background-color: #f8f9fa;
            }
        """)

        # **گروه ورودی فایل PDF**
        input_group = QGroupBox("فایل PDF ورودی")
        input_group.setAlignment(Qt.AlignLeft)
        input_group.setStyleSheet(group_style)
        input_layout = QVBoxLayout()
        input_layout.setSpacing(10)

        input_label = QLabel("فایل PDF ورودی:")
        input_label.setAlignment(Qt.AlignLeft)
        input_label.setStyleSheet(label_style)
        input_label.setToolTip("فایل PDF را که می‌خواهید قفل کنید، انتخاب کنید.")

        self.lock_input = QLineEdit()
        self.lock_input.setPlaceholderText("مسیر فایل PDF را انتخاب کنید")
        self.lock_input.setAlignment(Qt.AlignLeft)
        self.lock_input.setStyleSheet(input_style)
        self.lock_input.setToolTip("مسیر فایل PDF ورودی را وارد کنید یا از دکمه انتخاب فایل استفاده کنید.")

        browse_btn = QPushButton("انتخاب فایل")
        browse_btn.setStyleSheet(button_style)
        browse_btn.clicked.connect(lambda: self.lock_input.setText(QFileDialog.getOpenFileName(self, "انتخاب PDF", "", "فایل‌های PDF (*.pdf)")[0]))
        browse_btn.setToolTip("برای انتخاب فایل PDF، روی این دکمه کلیک کنید.")

        input_layout.addWidget(input_label)
        input_layout.addWidget(self.lock_input)
        input_layout.addWidget(browse_btn)
        input_group.setLayout(input_layout)

        # **گروه تنظیمات رمزگذاری**
        pwd_group = QGroupBox("تنظیمات رمزگذاری")
        pwd_group.setAlignment(Qt.AlignLeft)
        pwd_group.setStyleSheet(group_style)
        pwd_layout = QVBoxLayout()
        pwd_layout.setSpacing(10)

        read_pwd_label = QLabel("رمز برای باز کردن PDF:")
        read_pwd_label.setAlignment(Qt.AlignLeft)
        read_pwd_label.setStyleSheet(label_style)
        read_pwd_label.setToolTip("رمزی که برای باز کردن فایل PDF نیاز است.")

        self.lock_read_pwd = QLineEdit()
        self.lock_read_pwd.setEchoMode(QLineEdit.Password)
        self.lock_read_pwd.setPlaceholderText("رمز برای باز کردن PDF")
        self.lock_read_pwd.setAlignment(Qt.AlignLeft)
        self.lock_read_pwd.setStyleSheet(input_style)
        self.lock_read_pwd.setToolTip("رمزی را وارد کنید که برای باز کردن فایل PDF استفاده می‌شود.")

        self.lock_write_check = QCheckBox("استفاده از رمز متفاوت برای ویرایش")
        self.lock_write_check.setStyleSheet(checkbox_style)
        self.lock_write_check.setToolTip("برای استفاده از رمز جداگانه برای ویرایش فایل، این گزینه را فعال کنید.")

        write_pwd_label = QLabel("رمز برای ویرایش (اختیاری):")
        write_pwd_label.setAlignment(Qt.AlignLeft)
        write_pwd_label.setStyleSheet(label_style)
        write_pwd_label.setToolTip("رمزی که برای ویرایش فایل PDF نیاز است (اختیاری).")

        self.lock_write_pwd = QLineEdit()
        self.lock_write_pwd.setEchoMode(QLineEdit.Password)
        self.lock_write_pwd.setPlaceholderText("رمز برای ویرایش (اختیاری)")
        self.lock_write_pwd.setAlignment(Qt.AlignLeft)
        self.lock_write_pwd.setStyleSheet(input_style)
        self.lock_write_pwd.setEnabled(False)
        self.lock_write_pwd.setToolTip("رمزی را وارد کنید که برای ویرایش فایل PDF استفاده می‌شود (اگر متفاوت از رمز باز کردن باشد).")
        self.lock_write_check.stateChanged.connect(lambda state: self.lock_write_pwd.setEnabled(state == Qt.Checked))

        pwd_layout.addWidget(read_pwd_label)
        pwd_layout.addWidget(self.lock_read_pwd)
        pwd_layout.addWidget(self.lock_write_check)
        pwd_layout.addWidget(write_pwd_label)
        pwd_layout.addWidget(self.lock_write_pwd)
        pwd_group.setLayout(pwd_layout)

        # **گروه خروجی فایل**
        output_group = QGroupBox("فایل خروجی")
        output_group.setAlignment(Qt.AlignLeft)
        output_group.setStyleSheet(group_style)
        output_layout = QVBoxLayout()
        output_layout.setSpacing(10)

        output_label = QLabel("فایل خروجی:")
        output_label.setAlignment(Qt.AlignLeft)
        output_label.setStyleSheet(label_style)
        output_label.setToolTip("مسیر و نام فایل PDF قفل‌شده را مشخص کنید.")

        self.lock_output = QLineEdit()
        self.lock_output.setPlaceholderText("مسیر ذخیره فایل خروجی")
        self.lock_output.setAlignment(Qt.AlignLeft)
        self.lock_output.setStyleSheet(input_style)
        self.lock_output.setToolTip("مسیر و نام فایل PDF خروجی را وارد کنید یا از دکمه انتخاب مسیر استفاده کنید.")

        browse_output = QPushButton("انتخاب مسیر ذخیره")
        browse_output.setStyleSheet(button_style)
        browse_output.clicked.connect(lambda: self.lock_output.setText(QFileDialog.getSaveFileName(self, "ذخیره به عنوان", "", "فایل‌های PDF (*.pdf)")[0]))
        browse_output.setToolTip("برای انتخاب مسیر ذخیره فایل خروجی، روی این دکمه کلیک کنید.")

        output_layout.addWidget(output_label)
        output_layout.addWidget(self.lock_output)
        output_layout.addWidget(browse_output)
        output_group.setLayout(output_layout)

        # **دکمه قفل کردن**
        lock_btn = QPushButton("قفل کردن PDF")
        lock_btn.setStyleSheet(lock_button_style)
        lock_btn.clicked.connect(self.lock_pdf)
        lock_btn.setToolTip("برای قفل کردن فایل PDF با تنظیمات مشخص‌شده، روی این دکمه کلیک کنید.")

        # **افزودن گروه‌ها و دکمه به چیدمان اصلی**
        main_layout.addWidget(input_group)
        main_layout.addWidget(pwd_group)
        main_layout.addWidget(output_group)
        main_layout.addWidget(lock_btn, alignment=Qt.AlignCenter)

        widget.setLayout(main_layout)
        return widget

    def lock_pdf(self):
        input_file = self.lock_input.text().strip()
        output_file = self.lock_output.text().strip()
        read_pwd = self.lock_read_pwd.text().strip()
        write_pwd = self.lock_write_pwd.text().strip() if self.lock_write_check.isChecked() else read_pwd

        # بررسی ورودی‌ها
        if not all([input_file, output_file, read_pwd]):
            QMessageBox.warning(self, "هشدار", "لطفاً فیلدهای ورودی، خروجی و رمز خواندن را پر کنید.")
            return

        if not input_file.endswith('.pdf') or not output_file.endswith('.pdf'):
            QMessageBox.warning(self, "هشدار", "لطفاً فایل‌های PDF معتبر انتخاب کنید.")
            return

        try:
            with open(input_file, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                writer = PyPDF2.PdfWriter()
                for page in reader.pages:
                    writer.add_page(page)
                writer.encrypt(user_pwd=read_pwd, owner_pwd=write_pwd if write_pwd else None)
                with open(output_file, 'wb') as out:
                    writer.write(out)
            QMessageBox.information(self, "موفقیت", "PDF با موفقیت قفل شد.")
        except FileNotFoundError:
            QMessageBox.critical(self, "خطا", "فایل ورودی یافت نشد.")
        except PermissionError:
            QMessageBox.critical(self, "خطا", "دسترسی به فایل ممکن نیست. لطفاً مسیر دیگری انتخاب کنید.")
        except Exception as e:
            QMessageBox.critical(self, "خطا", f"قفل کردن PDF ناموفق بود: {str(e)}")


    def create_unlock_pdf_widget(self):
        # ایجاد ویجت اصلی
        widget = QWidget()
        widget.setLayoutDirection(Qt.RightToLeft)  # تنظیم جهت به راست‌به‌چپ
        main_layout = QVBoxLayout()
        main_layout.setAlignment(Qt.AlignCenter)  # تراز وسط برای زیبایی
        main_layout.setSpacing(20)  # فاصله بهینه بین گروه‌ها
        main_layout.setContentsMargins(30, 30, 30, 30)  # حاشیه‌های مدرن

        # استایل‌های مشترک
        input_style = """
            QLineEdit {
                padding: 10px;
                border-radius: 8px;
                border: 2px solid #e0e0e0;
                background-color: #ffffff;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QLineEdit:hover {
                border: 2px solid #007BFF;
                background-color: #f8faff;
            }
            QLineEdit:focus {
                border: 2px solid #0056b3;
                background-color: #ffffff;
            }
        """
        button_style = """
            QPushButton {
                padding: 10px;
                border-radius: 8px;
                background-color: #007BFF;
                color: white;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                font-weight: bold;
                border: none;
                min-width: 120px;
            }
            QPushButton:hover {
                background-color: #0056b3;
                border: 1px solid #004085;
            }
            QPushButton:pressed {
                background-color: #003d80;
            }
        """
        label_style = """
            QLabel {
                font-family: 'B Nazanin', 'Arial';
                font-size: 16px;
                font-weight: bold;
                color: #333333;
            }
        """
        group_style = """
            QGroupBox {
                border: 2px solid #e0e0e0;
                border-radius: 10px;
                padding: 15px;
                background-color: #ffffff;
                margin-top: 10px;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top Left;
                padding: 5px 10px;
                color: #333333;
                font-weight: bold;
            }
        """
        unlock_button_style = """
            QPushButton {
                padding: 12px;
                border-radius: 10px;
                background-color: #28A745;
                color: white;
                font-family: 'B Nazanin', 'Arial';
                font-size: 16px;
                font-weight: bold;
                border: none;
                min-width: 200px;
            }
            QPushButton:hover {
                background-color: #218838;
                border: 1px solid #1e7e34;
            }
            QPushButton:pressed {
                background-color: #1a6b2d;
            }
        """

        # تنظیم فونت و استایل کلی
        widget.setStyleSheet("""
            QWidget {
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                background-color: #f5f5f5;
            }
        """)

        # گروه فایل ورودی
        input_group = QGroupBox("فایل PDF ورودی")
        input_group.setAlignment(Qt.AlignLeft)
        input_group.setStyleSheet(group_style)
        input_layout = QVBoxLayout()
        input_layout.setSpacing(10)

        input_label = QLabel("فایل PDF ورودی:")
        input_label.setAlignment(Qt.AlignLeft)
        input_label.setStyleSheet(label_style)
        input_label.setToolTip("فایل PDF قفل‌شده را برای باز کردن قفل انتخاب کنید")

        self.unlock_input = QLineEdit()
        self.unlock_input.setPlaceholderText("مسیر فایل PDF را انتخاب کنید")
        self.unlock_input.setAlignment(Qt.AlignLeft)
        self.unlock_input.setStyleSheet(input_style)
        self.unlock_input.setToolTip("مسیر فایل PDF ورودی را وارد یا انتخاب کنید")

        browse_btn = QPushButton("انتخاب فایل")
        browse_btn.setStyleSheet(button_style)
        browse_btn.setToolTip("انتخاب فایل PDF")
        browse_btn.clicked.connect(lambda: self.unlock_input.setText(QFileDialog.getOpenFileName(self, "انتخاب PDF", "", "فایل‌های PDF (*.pdf)")[0]))

        input_layout.addWidget(input_label)
        input_layout.addWidget(self.unlock_input)
        input_layout.addWidget(browse_btn)
        input_group.setLayout(input_layout)

        # گروه رمز عبور
        pwd_group = QGroupBox("رمز عبور")
        pwd_group.setAlignment(Qt.AlignLeft)
        pwd_group.setStyleSheet(group_style)
        pwd_layout = QVBoxLayout()
        pwd_layout.setSpacing(10)

        pwd_label = QLabel("رمز عبور:")
        pwd_label.setAlignment(Qt.AlignLeft)
        pwd_label.setStyleSheet(label_style)
        pwd_label.setToolTip("رمز عبور فایل PDF را وارد کنید")

        self.unlock_pwd = QLineEdit()
        self.unlock_pwd.setPlaceholderText("رمز عبور فایل PDF را وارد کنید")
        self.unlock_pwd.setEchoMode(QLineEdit.Password)
        self.unlock_pwd.setAlignment(Qt.AlignLeft)
        self.unlock_pwd.setStyleSheet(input_style)
        self.unlock_pwd.setToolTip("رمز عبور فایل PDF را وارد کنید")

        pwd_layout.addWidget(pwd_label)
        pwd_layout.addWidget(self.unlock_pwd)
        pwd_group.setLayout(pwd_layout)

        # گروه فایل خروجی
        output_group = QGroupBox("فایل PDF خروجی")
        output_group.setAlignment(Qt.AlignLeft)
        output_group.setStyleSheet(group_style)
        output_layout = QVBoxLayout()
        output_layout.setSpacing(10)

        output_label = QLabel("فایل PDF خروجی:")
        output_label.setAlignment(Qt.AlignLeft)
        output_label.setStyleSheet(label_style)
        output_label.setToolTip("مسیر ذخیره فایل PDF باز شده را انتخاب کنید")

        self.unlock_output = QLineEdit()
        self.unlock_output.setPlaceholderText("مسیر ذخیره فایل PDF خروجی")
        self.unlock_output.setAlignment(Qt.AlignLeft)
        self.unlock_output.setStyleSheet(input_style)
        self.unlock_output.setToolTip("مسیر فایل PDF خروجی را وارد یا انتخاب کنید")

        browse_output = QPushButton("انتخاب مسیر")
        browse_output.setStyleSheet(button_style)
        browse_output.setToolTip("انتخاب مسیر ذخیره فایل خروجی")
        browse_output.clicked.connect(lambda: self.unlock_output.setText(QFileDialog.getSaveFileName(self, "ذخیره فایل PDF", "", "فایل‌های PDF (*.pdf)")[0]))

        output_layout.addWidget(output_label)
        output_layout.addWidget(self.unlock_output)
        output_layout.addWidget(browse_output)
        output_group.setLayout(output_layout)

        # دکمه باز کردن قفل
        unlock_btn = QPushButton("باز کردن قفل PDF")
        unlock_btn.setStyleSheet(unlock_button_style)
        unlock_btn.setToolTip("شروع فرآیند باز کردن قفل PDF")
        unlock_btn.clicked.connect(self.unlock_pdf)

        # افزودن گروه‌ها و دکمه به چیدمان اصلی
        main_layout.addWidget(input_group)
        main_layout.addWidget(pwd_group)
        main_layout.addWidget(output_group)
        main_layout.addWidget(unlock_btn, alignment=Qt.AlignCenter)
        main_layout.addStretch()  # فضای خالی برای زیبایی

        widget.setLayout(main_layout)
        return widget

    def unlock_pdf(self):
        input_file = self.unlock_input.text().strip()
        output_file = self.unlock_output.text().strip()
        password = self.unlock_pwd.text().strip()

        # بررسی ورودی‌ها
        if not all([input_file, output_file]):
            QMessageBox.warning(self, "هشدار", "لطفاً فیلدهای ورودی و خروجی را پر کنید.")
            return

        if not input_file.endswith('.pdf') or not output_file.endswith('.pdf'):
            QMessageBox.warning(self, "هشدار", "لطفاً فایل‌های PDF معتبر انتخاب کنید.")
            return

        try:
            with open(input_file, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                writer = PyPDF2.PdfWriter()
                
                if reader.is_encrypted:
                    if not password:
                        QMessageBox.warning(self, "هشدار", "لطفاً رمز عبور را وارد کنید.")
                        return
                    if not reader.decrypt(password):
                        QMessageBox.critical(self, "خطا", "رمز عبور نادرست است.")
                        return
                
                for page in reader.pages:
                    writer.add_page(page)
                
                with open(output_file, 'wb') as out:
                    writer.write(out)
            
            QMessageBox.information(self, "موفقیت", "قفل PDF با موفقیت باز شد.")
        except FileNotFoundError:
            QMessageBox.critical(self, "خطا", "فایل ورودی یافت نشد.")
        except PermissionError:
            QMessageBox.critical(self, "خطا", "دسترسی به فایل ممکن نیست. لطفاً مسیر دیگری انتخاب کنید.")
        except Exception as e:
            QMessageBox.critical(self, "خطا", f"باز کردن قفل PDF ناموفق بود: {str(e)}")






    def create_rotate_pages_widget(self):
        # ایجاد ویجت اصلی
        widget = QWidget()
        widget.setLayoutDirection(Qt.RightToLeft)  # تنظیم جهت به راست‌به‌چپ
        main_layout = QVBoxLayout()
        main_layout.setAlignment(Qt.AlignLeft)  # راست‌چین کردن چیدمان
        main_layout.setSpacing(20)  # فاصله مناسب بین گروه‌ها
        main_layout.setContentsMargins(30, 30, 30, 30)  # حاشیه‌های مدرن

        # استایل‌های مشترک
        input_style = """
            QLineEdit {
                padding: 8px;
                border-radius: 5px;
                border: 1px solid #ccc;
                background-color: #fff;
            }
            QLineEdit:hover {
                border: 1px solid #007BFF;
                background-color: #f0f8ff;
            }
        """
        button_style = """
            QPushButton {
                padding: 8px;
                border-radius: 5px;
                background-color: #007BFF;
                color: white;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #0056b3;
            }
            QPushButton:pressed {
                background-color: #003d80;
            }
        """
        label_style = "font-size: 14px; font-weight: bold; color: #333;"
        combo_style = """
            QComboBox {
                padding: 8px;
                border-radius: 5px;
                border: 1px solid #ccc;
                background-color: #fff;
            }
            QComboBox:hover {
                border: 1px solid #007BFF;
                background-color: #f0f8ff;
            }
        """
        group_style = """
            QGroupBox {
                border: 1px solid #ddd;
                border-radius: 5px;
                padding: 10px;
                background-color: #fff;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top Left;
                padding: 0 5px;
                color: #333;
                font-weight: bold;
            }
        """
        rotate_button_style = """
            QPushButton {
                padding: 10px;
                border-radius: 5px;
                background-color: #28A745;
                color: white;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #218838;
            }
            QPushButton:pressed {
                background-color: #1e7e34;
            }
        """

        # تنظیم فونت فارسی و استایل کلی
        widget.setStyleSheet("""
            QWidget {
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                background-color: #f8f9fa;
            }
        """)

        # گروه ورودی فایل PDF
        input_group = QGroupBox("فایل PDF ورودی")
        input_group.setAlignment(Qt.AlignLeft)
        input_group.setStyleSheet(group_style)
        input_layout = QVBoxLayout()
        input_layout.setSpacing(10)

        input_label = QLabel("فایل PDF ورودی:")
        input_label.setAlignment(Qt.AlignLeft)
        input_label.setStyleSheet(label_style)
        input_label.setToolTip("فایل PDF را که می‌خواهید صفحات آن را بچرخانید، انتخاب کنید.")

        self.rotate_input = QLineEdit()
        self.rotate_input.setPlaceholderText("مسیر فایل PDF را انتخاب کنید")
        self.rotate_input.setAlignment(Qt.AlignLeft)
        self.rotate_input.setStyleSheet(input_style)
        self.rotate_input.setToolTip("مسیر فایل PDF ورودی را وارد کنید یا از دکمه انتخاب فایل استفاده کنید.")

        browse_btn = QPushButton("انتخاب فایل")
        browse_btn.setStyleSheet(button_style)
        browse_btn.clicked.connect(lambda: self.rotate_input.setText(QFileDialog.getOpenFileName(self, "انتخاب PDF", "", "فایل‌های PDF (*.pdf)")[0]))
        browse_btn.setToolTip("برای انتخاب فایل PDF، روی این دکمه کلیک کنید.")

        input_layout.addWidget(input_label)
        input_layout.addWidget(self.rotate_input)
        input_layout.addWidget(browse_btn)
        input_group.setLayout(input_layout)

        # گروه تنظیمات چرخش
        rotate_group = QGroupBox("تنظیمات چرخش")
        rotate_group.setAlignment(Qt.AlignLeft)
        rotate_group.setStyleSheet(group_style)
        rotate_layout = QVBoxLayout()
        rotate_layout.setSpacing(10)

        pages_label = QLabel("صفحات (مثال: 1و2و3تا5 یا 1 2 4-6 یا خالی برای همه):")
        pages_label.setAlignment(Qt.AlignLeft)
        pages_label.setStyleSheet(label_style)
        pages_label.setToolTip("شماره صفحاتی که می‌خواهید بچرخانید را وارد کنید. برای همه صفحات، خالی بگذارید.")

        self.rotate_pages = QLineEdit()
        self.rotate_pages.setPlaceholderText("صفحات مورد نظر (مثال: 1و2و3تا5)")
        self.rotate_pages.setAlignment(Qt.AlignLeft)
        self.rotate_pages.setStyleSheet(input_style)
        self.rotate_pages.setToolTip("شماره صفحات را با جداکننده‌های مختلف یا بازه (با تا یا -) وارد کنید.")

        angle_label = QLabel("زاویه چرخش:")
        angle_label.setAlignment(Qt.AlignLeft)
        angle_label.setStyleSheet(label_style)
        angle_label.setToolTip("زاویه‌ای که می‌خواهید صفحات بچرخند را انتخاب کنید.")

        self.rotate_angle = QComboBox()
        self.rotate_angle.addItems(["0° (اولیه)", "90° ساعت‌گرد", "180° ساعت‌گرد", "270° ساعت‌گرد"])
        self.rotate_angle.setStyleSheet(combo_style)
        self.rotate_angle.setToolTip("زاویه چرخش را از میان گزینه‌های موجود انتخاب کنید.")

        rotate_layout.addWidget(pages_label)
        rotate_layout.addWidget(self.rotate_pages)
        rotate_layout.addWidget(angle_label)
        rotate_layout.addWidget(self.rotate_angle)
        rotate_group.setLayout(rotate_layout)

        # گروه خروجی فایل
        output_group = QGroupBox("فایل خروجی")
        output_group.setAlignment(Qt.AlignLeft)
        output_group.setStyleSheet(group_style)
        output_layout = QVBoxLayout()
        output_layout.setSpacing(10)

        output_label = QLabel("فایل خروجی:")
        output_label.setAlignment(Qt.AlignLeft)
        output_label.setStyleSheet(label_style)
        output_label.setToolTip("مسیر و نام فایل PDF خروجی را مشخص کنید.")

        self.rotate_output = QLineEdit()
        self.rotate_output.setPlaceholderText("مسیر ذخیره فایل خروجی")
        self.rotate_output.setAlignment(Qt.AlignLeft)
        self.rotate_output.setStyleSheet(input_style)
        self.rotate_output.setToolTip("مسیر و نام فایل PDF خروجی را وارد کنید یا از دکمه انتخاب مسیر استفاده کنید.")

        browse_output = QPushButton("انتخاب مسیر ذخیره")
        browse_output.setStyleSheet(button_style)
        browse_output.clicked.connect(lambda: self.rotate_output.setText(QFileDialog.getSaveFileName(self, "ذخیره فایل PDF", "", "فایل‌های PDF (*.pdf)")[0]))
        browse_output.setToolTip("برای انتخاب مسیر ذخیره فایل خروجی، روی این دکمه کلیک کنید.")

        output_layout.addWidget(output_label)
        output_layout.addWidget(self.rotate_output)
        output_layout.addWidget(browse_output)
        output_group.setLayout(output_layout)

        # دکمه چرخش صفحات
        rotate_btn = QPushButton("چرخش صفحات")
        rotate_btn.setStyleSheet(rotate_button_style)
        rotate_btn.clicked.connect(self.rotate_pages_func)
        rotate_btn.setToolTip("برای چرخش صفحات و ذخیره فایل جدید، روی این دکمه کلیک کنید.")

        # افزودن گروه‌ها و دکمه به چیدمان اصلی
        main_layout.addWidget(input_group)
        main_layout.addWidget(rotate_group)
        main_layout.addWidget(output_group)
        main_layout.addWidget(rotate_btn, alignment=Qt.AlignCenter)

        widget.setLayout(main_layout)
        return widget

    def rotate_pages_func(self):
        input_file = self.rotate_input.text().strip()
        output_file = self.rotate_output.text().strip()
        pages_text = self.rotate_pages.text().strip()
        angle = [0, 90, 180, 270][self.rotate_angle.currentIndex()]

        # بررسی ورودی‌ها
        if not input_file or not output_file:
            QMessageBox.warning(self, "هشدار", "لطفاً فیلدهای ورودی و خروجی را پر کنید.")
            return

        if not input_file.endswith('.pdf') or not output_file.endswith('.pdf'):
            QMessageBox.warning(self, "هشدار", "لطفاً فایل‌های PDF معتبر انتخاب کنید.")
            return

        try:
            # خواندن فایل PDF
            with open(input_file, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                total_pages = len(reader.pages)
                pages_to_rotate = self.parse_pages_input_r(pages_text, total_pages)

                # ایجاد فایل جدید با صفحات چرخانده‌شده
                writer = PyPDF2.PdfWriter()
                for i in range(total_pages):
                    page = reader.pages[i]
                    if i + 1 in pages_to_rotate:
                        page.rotate(angle)
                    writer.add_page(page)

                # ذخیره فایل خروجی
                with open(output_file, 'wb') as out:
                    writer.write(out)

            QMessageBox.information(self, "موفقیت", "صفحات با موفقیت چرخانده شدند.")
        except FileNotFoundError:
            QMessageBox.critical(self, "خطا", "فایل ورودی یافت نشد.")
        except PermissionError:
            QMessageBox.critical(self, "خطا", "دسترسی به فایل ممکن نیست. لطفاً مسیر دیگری انتخاب کنید.")
        except ValueError as e:
            QMessageBox.critical(self, "خطا", str(e))
        except Exception as e:
            QMessageBox.critical(self, "خطا", f"خطا در چرخش صفحات: {str(e)}")

    def parse_pages_input_r(self, input_str, total_pages):
        """
        ورودی کاربر را پارس می‌کند و لیستی از صفحات تکی یا بازه‌ها را برمی‌گرداند.
        پشتیبانی از ورودی‌هایی مثل: 1و2و3تا5 یا 1 2 4-6 یا 6تا4و1 2و3
        """
        if not input_str:
            return list(range(1, total_pages + 1))  # اگر خالی باشد، همه صفحات را برمی‌گرداند

        entries = []
        # جدا کردن بخش‌ها با استفاده از جداکننده‌های مختلف
        parts = re.split(r'[,\.\sو]+', input_str.strip())

        for part in parts:
            part = part.strip()
            if not part:
                continue
            # بررسی بازه‌ها (مثل 1تا3 یا 4-6)
            if 'تا' in part or '-' in part:
                range_parts = part.replace('تا', '-').split('-')
                if len(range_parts) == 2 and all(p.strip().isdigit() for p in range_parts):
                    start, end = map(int, range_parts)
                    # پشتیبانی از بازه‌های برعکس (مثل 5تا3)
                    start, end = min(start, end), max(start, end)
                    if 1 <= start <= total_pages and 1 <= end <= total_pages:
                        entries.extend(range(start, end + 1))
                    else:
                        raise ValueError(f"بازه {part} خارج از محدوده صفحات (1 تا {total_pages}) است.")
                else:
                    raise ValueError(f"فرمت بازه {part} نامعتبر است.")
            # بررسی صفحات تکی
            elif part.isdigit():
                page = int(part)
                if 1 <= page <= total_pages:
                    entries.append(page)
                else:
                    raise ValueError(f"صفحه {page} خارج از محدوده صفحات (1 تا {total_pages}) است.")
            else:
                raise ValueError(f"ورودی {part} نامعتبر است.")

        # حذف صفحات تکراری و مرتب‌سازی
        unique_pages = sorted(set(entries))
        return unique_pages



    def create_compress_pdf_widget(self):
        # ایجاد ویجت اصلی
        widget = QWidget()
        widget.setLayoutDirection(Qt.RightToLeft)  # راست‌چین کردن ویجت
        layout = QVBoxLayout()
        layout.setAlignment(Qt.AlignLeft)  # تنظیم alignment به راست
        layout.setSpacing(20)  # فاصله مناسب بین اجزا
        layout.setContentsMargins(20, 20, 20, 20)  # حاشیه‌های مدرن

                # استایل ورودی‌ها (QLineEdit)
        input_style = """
            QLineEdit {
                padding: 8px;
                border-radius: 5px;
                border: 1px solid #ccc;
                background-color: #fff;
            }
            QLineEdit:hover {
                border: 1px solid #007BFF;
                background-color: #f0f8ff;
            }
        """

        # استایل دکمه‌ها (QPushButton)
        button_style = """
            QPushButton {
                padding: 8px;
                border-radius: 5px;
                background-color: #007BFF;
                color: white;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #0056b3;
            }
            QPushButton:pressed {
                background-color: #003d80;
            }
        """

        # استایل برچسب‌ها (QLabel)
        label_style = """
            QLabel {
                font-size: 14px;
                font-weight: bold;
                color: #333;
            }
        """

        # استایل گروه‌ها (QGroupBox)
        group_style = """
            QGroupBox {
                border: 1px solid #ddd;
                border-radius: 5px;
                padding: 10px;
                background-color: #fafafa;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top Left;
                padding: 0 5px;
                color: #333;
                font-weight: bold;
            }
        """

        # استایل دکمه فشرده‌سازی
        compress_button_style = """
            QPushButton {
                padding: 10px;
                border-radius: 5px;
                background-color: #28A745;
                color: white;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #218838;
            }
            QPushButton:pressed {
                background-color: #1e7e34;
            }
        """

        # تنظیم فونت و استایل کلی
        widget.setStyleSheet("""
            QWidget {
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                background-color: #f8f9fa;
            }
        """)

        # گروه ورودی
        input_group = QGroupBox("فایل PDF ورودی")
        input_group.setAlignment(Qt.AlignLeft)
        input_group.setStyleSheet(group_style)
        input_layout = QVBoxLayout()
        input_layout.setSpacing(10)

        input_label = QLabel("فایل PDF ورودی:")
        input_label.setStyleSheet(label_style)
        input_label.setAlignment(Qt.AlignLeft)
        input_label.setToolTip("فایل PDF را برای فشرده‌سازی انتخاب کنید")

        self.compress_input = QLineEdit()
        self.compress_input.setPlaceholderText("مسیر فایل PDF را وارد کنید")
        self.compress_input.setAlignment(Qt.AlignLeft)
        self.compress_input.setStyleSheet(input_style)
        self.compress_input.setToolTip("مسیر فایل PDF را وارد یا از دکمه انتخاب فایل استفاده کنید")

        browse_btn = QPushButton("انتخاب فایل")
        browse_btn.setStyleSheet(button_style)
        browse_btn.clicked.connect(lambda: self.compress_input.setText(QFileDialog.getOpenFileName(self, "انتخاب PDF", "", "فایل‌های PDF (*.pdf)")[0]))
        browse_btn.setToolTip("برای انتخاب فایل PDF، روی این دکمه کلیک کنید")

        input_layout.addWidget(input_label)
        input_layout.addWidget(self.compress_input)
        input_layout.addWidget(browse_btn)
        input_group.setLayout(input_layout)

        # گروه خروجی
        output_group = QGroupBox("فایل PDF خروجی")
        output_group.setAlignment(Qt.AlignLeft)
        output_group.setStyleSheet(group_style)
        output_layout = QVBoxLayout()
        output_layout.setSpacing(10)

        output_label = QLabel("فایل PDF خروجی:")
        output_label.setStyleSheet(label_style)
        output_label.setAlignment(Qt.AlignLeft)
        output_label.setToolTip("مسیر ذخیره فایل PDF فشرده‌شده را مشخص کنید")

        self.compress_output = QLineEdit()
        self.compress_output.setPlaceholderText("مسیر ذخیره فایل PDF خروجی")
        self.compress_output.setAlignment(Qt.AlignLeft)
        self.compress_output.setStyleSheet(input_style)
        self.compress_output.setToolTip("مسیر فایل PDF خروجی را وارد یا از دکمه انتخاب مسیر استفاده کنید")

        browse_output = QPushButton("انتخاب مسیر")
        browse_output.setStyleSheet(button_style)
        browse_output.clicked.connect(lambda: self.compress_output.setText(QFileDialog.getSaveFileName(self, "ذخیره به‌عنوان", "", "فایل‌های PDF (*.pdf)")[0]))
        browse_output.setToolTip("برای انتخاب مسیر ذخیره فایل خروجی، روی این دکمه کلیک کنید")

        output_layout.addWidget(output_label)
        output_layout.addWidget(self.compress_output)
        output_layout.addWidget(browse_output)
        output_group.setLayout(output_layout)

        # دکمه فشرده‌سازی
        compress_btn = QPushButton("فشرده‌سازی PDF")
        compress_btn.setStyleSheet(compress_button_style)
        compress_btn.clicked.connect(self.compress_pdf)
        compress_btn.setToolTip("برای شروع فرآیند فشرده‌سازی، روی این دکمه کلیک کنید")

        # افزودن ویجت‌ها به چیدمان
        layout.addWidget(input_group)
        layout.addWidget(output_group)
        layout.addWidget(compress_btn, alignment=Qt.AlignCenter)
        layout.addStretch()  # فضای خالی در انتها برای زیبایی

        widget.setLayout(layout)
        return widget

    def compress_pdf(self):
        input_file = self.compress_input.text()
        output_file = self.compress_output.text()
        if not all([input_file, output_file]):
            QMessageBox.warning(self, "هشدار", "لطفاً تمام فیلدها را پر کنید.")
            return

        gs_path = self.find_ghostscript()
        if not gs_path:
            self.show_ghostscript_install_dialog()
            return

        try:
            gs_command = [
                gs_path, "-sDEVICE=pdfwrite", "-dCompatibilityLevel=1.4", "-dPDFSETTINGS=/screen",
                "-dNOPAUSE", "-dQUIET", "-dBATCH", "-dSAFER", f"-sOutputFile={output_file}", input_file
            ]
            subprocess.run(gs_command, check=True)
            QMessageBox.information(self, "موفقیت", "فایل PDF با موفقیت فشرده شد.")
        except subprocess.CalledProcessError as e:
            QMessageBox.critical(self, "خطا", f"فشرده‌سازی ناموفق بود: {e}")
        except Exception as e:
            QMessageBox.critical(self, "خطا", f"خطا در فشرده‌سازی: {e}")

    def find_ghostscript(self):
        possible_paths = [
            r"C:\Program Files\gs\gs10.05.0\bin\gswin64c.exe",
            r"C:\Program Files\gs\gs10.04.0\bin\gswin64c.exe",
            r"C:\Program Files (x86)\gs\gs10.05.0\bin\gswin32c.exe",
        ]
        for path in possible_paths:
            if os.path.exists(path):
                return path

        try:
            result = subprocess.run(["gs", "--version"], capture_output=True, text=True, check=True)
            if result.returncode == 0:
                return "gs"
        except (subprocess.CalledProcessError, FileNotFoundError):
            pass

        file_path, _ = QFileDialog.getOpenFileName(None, "انتخاب فایل اجرایی Ghostscript", "", 
                                                   "فایل‌های اجرایی (*.exe);;همه فایل‌ها (*)")
        return file_path if file_path else None

    def show_ghostscript_install_dialog(self):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Warning)
        msg.setText("Ghostscript یافت نشد.")
        msg.setInformativeText("برای فشرده‌سازی PDF، لطفاً Ghostscript را نصب کنید.")
        msg.setWindowTitle("نیاز به Ghostscript")
        msg.setLayoutDirection(Qt.RightToLeft)

        # استایل دکمه‌های پیام
        button_style = """
            QPushButton {
                padding: 8px; 
                border-radius: 5px; 
                background-color: #007BFF; 
                color: white;
            }
            QPushButton:hover {
                background-color: #0056b3;
            }
        """
        install_btn = QPushButton("نصب Ghostscript")
        install_btn.setStyleSheet(button_style)
        install_btn.clicked.connect(self.install_ghostscript)
        msg.addButton(install_btn, QMessageBox.ActionRole)
        ok_btn = msg.addButton(QMessageBox.Ok)
        ok_btn.setStyleSheet(button_style)

        msg.exec_()

    def install_ghostscript(self):
        if getattr(sys, 'frozen', False):
            base_path = os.path.dirname(sys.executable)
        else:
            base_path = os.path.dirname(os.path.abspath(__file__))

        installer_path = os.path.join(base_path, "gs10050w64.exe")

        if os.path.exists(installer_path):
            try:
                ctypes.windll.shell32.ShellExecuteW(None, "runas", installer_path, None, None, 1)
            except Exception as e:
                QMessageBox.critical(self, "خطا", f"خطا در اجرای نصب: {e}")
        else:
            QMessageBox.critical(self, "خطا", "فایل نصب Ghostscript یافت نشد. لطفاً آن را به‌صورت دستی نصب کنید.")

    def check_ghostscript(self):
        gs_path = self.find_ghostscript()
        if gs_path and gs_path != "gs":
            gs_dir = os.path.dirname(gs_path)
            current_path = os.environ.get('PATH', '')
            if gs_dir not in current_path.split(';'):
                try:
                    new_path = f"{current_path};{gs_dir}"
                    subprocess.run(f'setx PATH "{new_path}"', check=True, shell=True)
                except Exception as e:
                    QMessageBox.critical(self, "خطا", f"خطا در افزودن Ghostscript به PATH: {e}")

                

    def create_extract_text_widget(self):
        # ایجاد ویجت اصلی
        widget = QWidget()
        layout = QVBoxLayout()

        # استایل‌های مشترک
        input_style = """
            QLineEdit { padding: 8px; border-radius: 5px; border: 1px solid #ccc; background-color: #fff; }
            QLineEdit:hover { border: 1px solid #007BFF; background-color: #f8faff; }
        """
        button_style = """
            QPushButton { padding: 8px; border-radius: 5px; background-color: #007BFF; color: white; font-weight: bold; }
            QPushButton:hover { background-color: #0056b3; }
        """
        extract_button_style = """
            QPushButton { padding: 10px; border-radius: 5px; background-color: #28A745; color: white; font-weight: bold; }
            QPushButton:hover { background-color: #218838; }
        """
        label_style = """
            QLabel { font-size: 14px; font-weight: bold; color: #333; }
        """
        group_style = """
            QGroupBox { border: 1px solid #ddd; border-radius: 5px; padding: 10px; } 
            QGroupBox::title { subcontrol-origin: margin; subcontrol-position: top Left; padding: 0 5px; color: #333; font-weight: bold; }
        """
        progress_style = """
            QProgressBar { border: 1px solid #ccc; border-radius: 5px; text-align: center; background-color: #f5f5f5; }
            QProgressBar::chunk { background-color: #28A745; border-radius: 5px; }
        """

        # تنظیم فونت و استایل کلی
        widget.setStyleSheet("QWidget { font-family: 'B Nazanin', 'Arial'; font-size: 14px; background-color: #f5f5f5; }")
        widget.setLayoutDirection(Qt.RightToLeft)

        # گروه‌بندی ویجت‌ها
        group_box = QGroupBox("استخراج متن و تصاویر از PDF")
        group_box.setStyleSheet(group_style)
        group_layout = QVBoxLayout()

        # فیلد ورودی PDF
        self.text_input = QLineEdit()
        self.text_input.setStyleSheet(input_style)
        self.text_input.setPlaceholderText("مسیر فایل PDF را وارد کنید")
        self.text_input.setToolTip("فایل PDF را برای استخراج متن و تصاویر انتخاب کنید")
        browse_btn = QPushButton("انتخاب فایل")
        browse_btn.setStyleSheet(button_style)
        browse_btn.setToolTip("برای انتخاب فایل PDF کلیک کنید")
        browse_btn.clicked.connect(lambda: self.text_input.setText(QFileDialog.getOpenFileName(self, "انتخاب PDF", "", "فایل‌های PDF (*.pdf)")[0]))

        # فیلد صفحات
        self.text_pages = QLineEdit()
        self.text_pages.setStyleSheet(input_style)
        self.text_pages.setPlaceholderText("صفحات (مثال: 4,8,6-24,42 یا خالی برای همه)")
        self.text_pages.setToolTip("شماره صفحاتی که می‌خواهید استخراج شود (مثال: 1و2و3تا5)")

        # فیلد خروجی
        self.text_output = QLineEdit()
        self.text_output.setStyleSheet(input_style)
        self.text_output.setPlaceholderText("مسیر فایل خروجی را وارد کنید")
        self.text_output.setToolTip("مسیر و نام فایل خروجی را مشخص کنید")
        browse_output = QPushButton("انتخاب مقصد")
        browse_output.setStyleSheet(button_style)
        browse_output.setToolTip("برای انتخاب مسیر فایل خروجی کلیک کنید")
        browse_output.clicked.connect(self.browse_output_file_e)

        # دکمه استخراج
        extract_btn = QPushButton("استخراج متن و تصاویر")
        extract_btn.setStyleSheet(extract_button_style)
        extract_btn.setToolTip("برای شروع فرآیند استخراج کلیک کنید")
        extract_btn.clicked.connect(self.extract_text)

        # نوار پیشرفت
        self.progress_bar = QProgressBar()
        self.progress_bar.setStyleSheet(progress_style)
        self.progress_bar.setValue(0)
        self.progress_bar.setToolTip("پیشرفت فرآیند استخراج را نشان می‌دهد")

        # افزودن ویجت‌ها به چیدمان گروه
        group_layout.addWidget(QLabel("فایل PDF ورودی:", styleSheet=label_style))
        group_layout.addWidget(self.text_input)
        group_layout.addWidget(browse_btn)
        group_layout.addWidget(QLabel("صفحات:", styleSheet=label_style))
        group_layout.addWidget(self.text_pages)
        group_layout.addWidget(QLabel("فایل خروجی:", styleSheet=label_style))
        group_layout.addWidget(self.text_output)
        group_layout.addWidget(browse_output)
        group_layout.addWidget(extract_btn)
        group_layout.addWidget(QLabel("پیشرفت:", styleSheet=label_style))
        group_layout.addWidget(self.progress_bar)

        group_box.setLayout(group_layout)
        layout.addWidget(group_box)
        layout.setSpacing(20)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.addStretch()
        widget.setLayout(layout)
        return widget

    def browse_output_file_e(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "ذخیره به‌عنوان", "", "فایل‌های ورد (*.docx)")
        if file_path:
            if not file_path.endswith('.docx'):
                file_path += '.docx'
            self.text_output.setText(file_path)
        else:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Warning)
            msg.setWindowTitle("هشدار")
            msg.setText("لطفاً یک مسیر معتبر برای فایل خروجی انتخاب کنید.")
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()

    def extract_text(self):
        input_file = self.text_input.text()
        output_file = self.text_output.text()
        pages_text = self.text_pages.text()

        if not all([input_file, output_file]):
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Warning)
            msg.setWindowTitle("هشدار")
            msg.setText("لطفاً فیلدهای ورودی و خروجی را پر کنید.")
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
            return

        if os.path.exists(output_file):
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Question)
            msg.setWindowTitle("هشدار")
            msg.setText("فایل خروجی وجود دارد. بازنویسی شود؟")
            msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
            msg.setDefaultButton(QMessageBox.No)
            msg.setLayoutDirection(Qt.RightToLeft)
            result = msg.exec_()
            if result == QMessageBox.No:
                return

        try:
            doc_pdf = fitz.open(input_file)
            total_pages = doc_pdf.page_count
            pages = self.parse_pages_input_e(pages_text, total_pages) if pages_text else list(range(1, total_pages + 1))

            doc_word = Document()
            section = doc_word.sections[0]
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.page_width = Cm(29.7)  # A4 افقی
            section.page_height = Cm(21.0)

            self.progress_bar.setMaximum(len(pages))
            self.progress_bar.setValue(0)

            for idx, page_num in enumerate(pages):
                page = doc_pdf.load_page(page_num - 1)
                text_instances = page.get_text("dict")["blocks"]
                doc_word.add_paragraph(f"(صفحه {page_num}):", style='Heading 1')

                # استخراج و افزودن تصاویر
                images = page.get_images(full=True)
                for img_index, img in enumerate(images):
                    xref = img[0]
                    base_image = doc_pdf.extract_image(xref)
                    image_bytes = base_image["image"]
                    image = Image.open(io.BytesIO(image_bytes))
                    image_stream = io.BytesIO()
                    image.save(image_stream, format="PNG")
                    p = doc_word.add_paragraph()
                    run = p.add_run()
                    run.add_picture(image_stream, width=Cm(base_image["width"] / 50))

                # استخراج متن
                for block in text_instances:
                    if block['type'] == 0:  # بلوک متنی
                        for line in block['lines']:
                            line_text = ""
                            font_size = 12
                            is_bold = False
                            color = (0, 0, 0)
                            for span in line['spans']:
                                text = span['text'].strip()
                                if text:
                                    text = unicodedata.normalize('NFKC', text)
                                    line_text += text + " "
                                    font_size = span.get('size', 12)
                                    flags = span.get('flags', 0)
                                    is_bold = bool(flags & 2)
                                    span_color = span.get('color', 0)
                                    color = self.hex_to_rgb(span_color)

                            if line_text.strip():
                                is_persian = re.search(r'[\u0600-\u06FF]', line_text)
                                p = doc_word.add_paragraph()
                                p.paragraph_format.line_spacing = 1.15
                                p.paragraph_format.space_after = Pt(6)
                                bbox = line['bbox']
                                indent_left = Cm(bbox[0] / 30)
                                p.paragraph_format.left_indent = indent_left

                                if is_persian:
                                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    run = p.add_run(line_text.strip())
                                    run.font.name = 'B Nazanin'
                                    self.set_run_direction(run, 'rtl')
                                else:
                                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    run = p.add_run(line_text.strip())
                                    run.font.name = 'Times New Roman'
                                    self.set_run_direction(run, 'ltr')

                                run.font.size = Pt(round(font_size))
                                run.bold = is_bold
                                run.font.color.rgb = RGBColor(*color)

                doc_word.add_page_break()
                self.progress_bar.setValue(idx + 1)

            doc_word.save(output_file)
            doc_pdf.close()
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Information)
            msg.setWindowTitle("موفقیت")
            msg.setText("متن و تصاویر با موفقیت استخراج شد.")
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()

        except FileNotFoundError:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Critical)
            msg.setWindowTitle("خطا")
            msg.setText("فایل PDF یافت نشد.")
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
        except ValueError as ve:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Critical)
            msg.setWindowTitle("خطا")
            msg.setText(f"ورودی نامعتبر: {str(ve)}")
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
        except Exception as e:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Critical)
            msg.setWindowTitle("خطا")
            msg.setText(f"خطا در استخراج: {str(e)}")
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()

    def hex_to_rgb(self, hex_color):
        """تبدیل رنگ از فرمت hex به RGB"""
        hex_color = hex_color & 0xFFFFFF
        r = (hex_color >> 16) & 255
        g = (hex_color >> 8) & 255
        b = hex_color & 255
        return (r, g, b)

    def set_run_direction(self, run, direction):
        """تنظیم جهت متن در Word"""
        run.element.set(qn('w:bidi'), '1' if direction == 'rtl' else '0')

    def parse_pages_input_e(self, input_str, total_pages):
        """
        پارس ورودی صفحات و پشتیبانی از فرمت‌هایی مثل: 1و2و3تا5، 1 2 3و4تا6، 6تا4و1 2و3
        خروجی: لیستی از شماره صفحات بدون تکرار
        """
        pages = set()  # استفاده از مجموعه برای حذف تکرار
        parts = re.split(r'[,\.\sو]+', input_str.strip())

        for part in parts:
            part = part.strip()
            if not part:
                continue
            if 'تا' in part or '-' in part:
                range_parts = part.replace('تا', '-').split('-')
                if len(range_parts) == 2 and all(p.strip().isdigit() for p in range_parts):
                    start, end = map(int, range_parts)
                    start, end = min(start, end), max(start, end)  # پشتیبانی از بازه‌های برعکس
                    if 1 <= start <= total_pages and 1 <= end <= total_pages:
                        pages.update(range(start, end + 1))
                    else:
                        raise ValueError(f"بازه {part} خارج از محدوده صفحات (1 تا {total_pages}) است.")
                else:
                    raise ValueError(f"فرمت بازه {part} نامعتبر است.")
            elif part.isdigit():
                page = int(part)
                if 1 <= page <= total_pages:
                    pages.add(page)
                else:
                    raise ValueError(f"صفحه {page} خارج از محدوده صفحات (1 تا {total_pages}) است.")
            else:
                raise ValueError(f"ورودی {part} نامعتبر است.")

        return sorted(pages)





    def create_extract_images_widget(self):
        # ویجت اصلی
        widget = QWidget()
        widget.setLayoutDirection(Qt.RightToLeft)  # تنظیم جهت به راست‌به‌چپ برای زبان فارسی
        main_layout = QVBoxLayout()
        main_layout.setAlignment(Qt.AlignLeft)
        main_layout.setSpacing(20)
        main_layout.setContentsMargins(30, 30, 30, 30)

        # استایل‌های مشترک
        input_style = """
            QLineEdit {
                padding: 10px;
                border-radius: 8px;
                border: 2px solid #e0e0e0;
                background-color: #ffffff;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QLineEdit:hover {
                border: 2px solid #007BFF;
                background-color: #f8faff;
            }
            QLineEdit:focus {
                border: 2px solid #0056b3;
                background-color: #ffffff;
            }
        """
        button_style = """
            QPushButton {
                padding: 10px;
                border-radius: 8px;
                background-color: #007BFF;
                color: white;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                font-weight: bold;
                border: none;
                min-width: 120px;
            }
            QPushButton:hover {
                background-color: #0056b3;
                border: 1px solid #004085;
            }
            QPushButton:pressed {
                background-color: #003d80;
            }
        """
        extract_button_style = """
            QPushButton {
                padding: 12px;
                border-radius: 10px;
                background-color: #28A745;
                color: white;
                font-family: 'B Nazanin', 'Arial';
                font-size: 16px;
                font-weight: bold;
                border: none;
                min-width: 200px;
            }
            QPushButton:hover {
                background-color: #218838;
                border: 1px solid #1e7e34;
            }
            QPushButton:pressed {
                background-color: #1a6b2d;
            }
        """
        label_style = """
            QLabel {
                font-family: 'B Nazanin', 'Arial';
                font-size: 16px;
                font-weight: bold;
                color: #333333;
            }
        """
        group_style = """
            QGroupBox {
                border: 2px solid #e0e0e0;
                border-radius: 10px;
                padding: 15px;
                background-color: #ffffff;
                margin-top: 10px;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top left;
                padding: 5px 10px;
                color: #333333;
                font-weight: bold;
            }
        """

        # تنظیم فونت و استایل کلی
        widget.setStyleSheet("""
            QWidget {
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                background-color: #f5f5f5;
            }
        """)

        # گروه استخراج تصاویر
        group_box = QGroupBox("استخراج تصاویر از PDF")
        group_box.setAlignment(Qt.AlignLeft)
        group_box.setStyleSheet(group_style)
        group_layout = QVBoxLayout()
        group_layout.setSpacing(10)

        # فیلد ورودی PDF
        input_label = QLabel("فایل PDF ورودی:")
        input_label.setStyleSheet(label_style)
        input_label.setAlignment(Qt.AlignLeft)
        input_label.setToolTip("فایل PDF را برای استخراج تصاویر انتخاب کنید")

        self.images_input = QLineEdit()
        self.images_input.setPlaceholderText("مسیر فایل PDF را انتخاب کنید")
        self.images_input.setAlignment(Qt.AlignLeft)
        self.images_input.setStyleSheet(input_style)
        self.images_input.setToolTip("مسیر فایل PDF ورودی را وارد یا انتخاب کنید")

        browse_btn = QPushButton("انتخاب فایل")
        browse_btn.setStyleSheet(button_style)
        browse_btn.setToolTip("انتخاب فایل PDF")
        browse_btn.clicked.connect(lambda: self.images_input.setText(QFileDialog.getOpenFileName(self, "انتخاب PDF", "", "فایل‌های PDF (*.pdf)")[0]))

        # فیلد صفحات
        pages_label = QLabel("صفحات (مثال: 1و2و3و4تا5، خالی برای همه):")
        pages_label.setStyleSheet(label_style)
        pages_label.setAlignment(Qt.AlignLeft)
        pages_label.setToolTip("شماره صفحاتی که می‌خواهید تصاویر از آن‌ها استخراج شود را وارد کنید")

        self.images_pages = QLineEdit()
        self.images_pages.setPlaceholderText("صفحات مورد نظر (مثال: 1و2و3و4تا5)")
        self.images_pages.setAlignment(Qt.AlignLeft)
        self.images_pages.setStyleSheet(input_style)
        self.images_pages.setToolTip("شماره صفحات را وارد کنید یا برای همه صفحات خالی بگذارید")

        # فیلد دایرکتوری خروجی
        output_label = QLabel("پوشه خروجی:")
        output_label.setStyleSheet(label_style)
        output_label.setAlignment(Qt.AlignLeft)
        output_label.setToolTip("پوشه‌ای که تصاویر استخراج‌شده در آن ذخیره می‌شوند")

        self.images_output_dir = QLineEdit()
        self.images_output_dir.setPlaceholderText("مسیر پوشه خروجی را انتخاب کنید")
        self.images_output_dir.setAlignment(Qt.AlignLeft)
        self.images_output_dir.setStyleSheet(input_style)
        self.images_output_dir.setToolTip("مسیر پوشه‌ای که تصاویر در آن ذخیره می‌شوند")

        browse_dir = QPushButton("انتخاب پوشه")
        browse_dir.setStyleSheet(button_style)
        browse_dir.setToolTip("انتخاب پوشه خروجی")
        browse_dir.clicked.connect(lambda: self.images_output_dir.setText(QFileDialog.getExistingDirectory(self, "انتخاب پوشه خروجی")))

        # دکمه استخراج
        extract_btn = QPushButton("استخراج تصاویر")
        extract_btn.setStyleSheet(extract_button_style)
        extract_btn.setToolTip("شروع فرآیند استخراج تصاویر")
        extract_btn.clicked.connect(self.extract_images)

        # افزودن ویجت‌ها به چیدمان گروه
        group_layout.addWidget(input_label)
        group_layout.addWidget(self.images_input)
        group_layout.addWidget(browse_btn)
        group_layout.addWidget(pages_label)
        group_layout.addWidget(self.images_pages)
        group_layout.addWidget(output_label)
        group_layout.addWidget(self.images_output_dir)
        group_layout.addWidget(browse_dir)
        group_layout.addWidget(extract_btn)

        group_box.setLayout(group_layout)

        # افزودن گروه به چیدمان اصلی
        main_layout.addWidget(group_box)
        main_layout.addStretch()

        widget.setLayout(main_layout)
        return widget

    def extract_images(self):
        input_file = self.images_input.text()
        output_dir = self.images_output_dir.text()
        pages_text = self.images_pages.text()

        # اعتبارسنجی ورودی‌ها
        if not all([input_file, output_dir]):
            msg = QMessageBox()
            msg.setWindowTitle("هشدار")
            msg.setText("لطفاً فیلدهای ورودی و خروجی را پر کنید.")
            msg.setIcon(QMessageBox.Warning)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
            return

        # بررسی دسترسی به پوشه خروجی
        if not os.access(output_dir, os.W_OK):
            msg = QMessageBox()
            msg.setWindowTitle("خطا")
            msg.setText("پوشه خروجی قابل نوشتن نیست. لطفاً پوشه دیگری انتخاب کنید.")
            msg.setIcon(QMessageBox.Critical)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
            return

        try:
            with pdfplumber.open(input_file) as pdf:
                total_pages = len(pdf.pages)
                # اگر فیلد صفحات خالی باشد، همه صفحات انتخاب می‌شوند
                pages = self.parse_pages_input_i(pages_text, total_pages) if pages_text else range(1, total_pages + 1)

                for i in pages:
                    page = pdf.pages[i - 1]
                    page_bbox = (0, 0, page.width, page.height)  # جعبه مرزی صفحه
                    for j, img in enumerate(page.images):
                        # جعبه مرزی تصویر
                        img_bbox = (img['x0'], img['top'], img['x1'], img['bottom'])
                        # محدود کردن جعبه مرزی به صفحه
                        intersected_bbox = (
                            max(img_bbox[0], page_bbox[0]),  # x0
                            max(img_bbox[1], page_bbox[1]),  # top
                            min(img_bbox[2], page_bbox[2]),  # x1
                            min(img_bbox[3], page_bbox[3])   # bottom
                        )
                        # بررسی اینکه آیا تصویر در محدوده صفحه است
                        if intersected_bbox[0] < intersected_bbox[2] and intersected_bbox[1] < intersected_bbox[3]:
                            cropped = page.crop(intersected_bbox).to_image(resolution=300)
                            output_path = os.path.join(output_dir, f"صفحه{i}_تصویر{j}.png")
                            cropped.save(output_path)
                        else:
                            print(f"تصویر {j} در صفحه {i} خارج از محدوده است و نادیده گرفته شد.")

            msg = QMessageBox()
            msg.setWindowTitle("موفقیت")
            msg.setText("تصاویر با موفقیت استخراج شدند.")
            msg.setIcon(QMessageBox.Information)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()

        except FileNotFoundError:
            msg = QMessageBox()
            msg.setWindowTitle("خطا")
            msg.setText("فایل PDF یافت نشد.")
            msg.setIcon(QMessageBox.Critical)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
        except ValueError as ve:
            msg = QMessageBox()
            msg.setWindowTitle("خطا")
            msg.setText(f"ورودی نامعتبر: {str(ve)}")
            msg.setIcon(QMessageBox.Critical)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
        except Exception as e:
            msg = QMessageBox()
            msg.setWindowTitle("خطا")
            msg.setText(f"خطا در استخراج تصاویر: {str(e)}")
            msg.setIcon(QMessageBox.Critical)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()

    def parse_pages_input_i(self, input_str, total_pages):
        """
        پارس کردن ورودی صفحات با پشتیبانی از فرمت‌های:
        - 1و2و3و4تا5
        - 1 2 3و4تا6
        - 6تا4و1 2و3
        خروجی: لیستی از شماره صفحات بدون تکرار
        """
        pages = set()
        parts = re.split(r'[,\.\sو]+', input_str.strip())

        for part in parts:
            part = part.strip()
            if not part:
                continue
            if 'تا' in part or '-' in part:
                range_parts = part.replace('تا', '-').split('-')
                if len(range_parts) == 2 and all(p.strip().isdigit() for p in range_parts):
                    start, end = map(int, range_parts)
                    start, end = min(start, end), max(start, end)  # پشتیبانی از بازه معکوس مثل 6تا4
                    if 1 <= start <= total_pages and 1 <= end <= total_pages:
                        pages.update(range(start, end + 1))
                    else:
                        raise ValueError(f"بازه {part} خارج از محدوده صفحات (1 تا {total_pages}) است.")
                else:
                    raise ValueError(f"فرمت بازه {part} نامعتبر است.")
            elif part.isdigit():
                page = int(part)
                if 1 <= page <= total_pages:
                    pages.add(page)
                else:
                    raise ValueError(f"صفحه {page} خارج از محدوده صفحات (1 تا {total_pages}) است.")
            else:
                raise ValueError(f"ورودی {part} نامعتبر است.")

        return sorted(pages)





    def create_image_to_pdf_widget(self):
        # ایجاد ویجت اصلی
        widget = QWidget()
        widget.setLayoutDirection(Qt.RightToLeft)  # چپ‌چین کردن ویجت
        main_layout = QVBoxLayout()
        main_layout.setAlignment(Qt.AlignLeft)
        main_layout.setSpacing(20)  # فاصله بهینه بین گروه‌ها
        main_layout.setContentsMargins(30, 30, 30, 30)  # حاشیه‌های مدرن

        # استایل‌های مشترک
        input_style = """
            QLineEdit {
                padding: 10px;
                border-radius: 8px;
                border: 2px solid #e0e0e0;
                background-color: #ffffff;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QLineEdit:hover {
                border: 2px solid #007BFF;
                background-color: #f8faff;
            }
            QLineEdit:focus {
                border: 2px solid #0056b3;
                background-color: #ffffff;
            }
        """
        button_style = """
            QPushButton {
                padding: 10px;
                border-radius: 8px;
                background-color: #007BFF;
                color: white;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                font-weight: bold;
                border: none;
                min-width: 120px;
            }
            QPushButton:hover {
                background-color: #0056b3;
                border: 1px solid #004085;
            }
            QPushButton:pressed {
                background-color: #003d80;
            }
        """
        convert_button_style = """
            QPushButton {
                padding: 12px;
                border-radius: 10px;
                background-color: #28A745;
                color: white;
                font-family: 'B Nazanin', 'Arial';
                font-size: 16px;
                font-weight: bold;
                border: none;
                min-width: 200px;
            }
            QPushButton:hover {
                background-color: #218838;
                border: 1px solid #1e7e34;
            }
            QPushButton:pressed {
                background-color: #1a6b2d;
            }
        """
        label_style = """
            QLabel {
                font-family: 'B Nazanin', 'Arial';
                font-size: 16px;
                font-weight: bold;
                color: #333333;
            }
        """
        group_style = """
            QGroupBox {
                border: 2px solid #e0e0e0;
                border-radius: 10px;
                padding: 15px;
                background-color: #ffffff;
                margin-top: 10px;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top left;
                padding: 5px 10px;
                color: #333333;
                font-weight: bold;
            }
        """
        list_style = """
            QListWidget {
                padding: 10px;
                border-radius: 8px;
                border: 2px solid #e0e0e0;
                background-color: #ffffff;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QListWidget::item:hover {
                background-color: #f0f8ff;
            }
            QListWidget::item:selected {
                background-color: #007BFF;
                color: white;
            }
        """
        checkbox_style = """
            QCheckBox {
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                color: #333333;
                padding: 5px;
            }
            QCheckBox:hover {
                color: #007BFF;
            }
        """
        spinbox_style = """
            QSpinBox {
                padding: 10px;
                border-radius: 8px;
                border: 2px solid #e0e0e0;
                background-color: #ffffff;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QSpinBox:hover {
                border: 2px solid #007BFF;
                background-color: #f8faff;
            }
            QSpinBox:focus {
                border: 2px solid #0056b3;
                background-color: #ffffff;
            }
        """

        # تنظیم فونت و استایل کلی
        widget.setStyleSheet("""
            QWidget {
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                background-color: #f5f5f5;
            }
        """)

        # گروه تصاویر
        images_group = QGroupBox("مدیریت تصاویر")
        images_group.setAlignment(Qt.AlignLeft)
        images_group.setStyleSheet(group_style)
        images_layout = QVBoxLayout()
        images_layout.setSpacing(10)

        images_label = QLabel("تصاویر ورودی:")
        images_label.setStyleSheet(label_style)
        images_label.setAlignment(Qt.AlignLeft)
        images_label.setToolTip("تصاویر را برای تبدیل به PDF انتخاب کنید")

        self.image_list = QListWidget()
        self.image_list.setLayoutDirection(Qt.RightToLeft)
        self.image_list.setStyleSheet(list_style)
        self.image_list.setMinimumHeight(150)
        self.image_list.setToolTip("لیست تصاویر انتخاب‌شده برای تبدیل به PDF")

        buttons_layout = QHBoxLayout()
        buttons_layout.setSpacing(10)
        add_btn = QPushButton("افزودن تصاویر")
        add_btn.setStyleSheet(button_style)
        add_btn.setToolTip("افزودن تصاویر جدید به لیست")
        add_btn.clicked.connect(lambda: self.image_list.addItems(QFileDialog.getOpenFileNames(self, "انتخاب تصاویر", "", "فایل‌های تصویری (*.png *.jpg *.jpeg)")[0]))

        remove_btn = QPushButton("حذف انتخاب‌شده‌ها")
        remove_btn.setStyleSheet(button_style)
        remove_btn.setToolTip("حذف تصاویر انتخاب‌شده از لیست")
        remove_btn.clicked.connect(lambda: [self.image_list.takeItem(self.image_list.row(item)) for item in self.image_list.selectedItems()])

        buttons_layout.addWidget(add_btn)
        buttons_layout.addWidget(remove_btn)
        buttons_layout.addStretch()

        images_layout.addWidget(images_label)
        images_layout.addWidget(self.image_list)
        images_layout.addLayout(buttons_layout)
        images_group.setLayout(images_layout)

        # گروه تنظیمات صفحه
        settings_group = QGroupBox("تنظیمات صفحه")
        settings_group.setAlignment(Qt.AlignLeft)
        settings_group.setStyleSheet(group_style)
        settings_layout = QGridLayout()
        settings_layout.setSpacing(10)

        self.image_one_per_page = QCheckBox("یک تصویر در هر صفحه")
        self.image_one_per_page.setStyleSheet(checkbox_style)
        self.image_one_per_page.setToolTip("هر تصویر در یک صفحه جداگانه قرار گیرد")

        self.image_center = QCheckBox("مرکز صفحه")
        self.image_center.setStyleSheet(checkbox_style)
        self.image_center.setToolTip("تصویر در مرکز صفحه قرار گیرد")

        self.image_width = QSpinBox()
        self.image_width.setRange(0, 1000)
        self.image_width.setStyleSheet(spinbox_style)
        self.image_width.setToolTip("عرض تصویر در PDF (0 برای اندازه اصلی)")

        self.image_height = QSpinBox()
        self.image_height.setRange(0, 1000)
        self.image_height.setStyleSheet(spinbox_style)
        self.image_height.setToolTip("ارتفاع تصویر در PDF (0 برای اندازه اصلی)")

        self.image_x_offset = QSpinBox()
        self.image_x_offset.setRange(-500, 500)
        self.image_x_offset.setStyleSheet(spinbox_style)
        self.image_x_offset.setToolTip("انحراف افقی تصویر")

        self.image_y_offset = QSpinBox()
        self.image_y_offset.setRange(-500, 500)
        self.image_y_offset.setStyleSheet(spinbox_style)
        self.image_y_offset.setToolTip("انحراف عمودی تصویر")

        width_label = QLabel("عرض تصویر (0 برای اصلی):")
        width_label.setStyleSheet(label_style)
        height_label = QLabel("ارتفاع تصویر (0 برای اصلی):")
        height_label.setStyleSheet(label_style)
        x_offset_label = QLabel("انحراف افقی:")
        x_offset_label.setStyleSheet(label_style)
        y_offset_label = QLabel("انحراف عمودی:")
        y_offset_label.setStyleSheet(label_style)

        settings_layout.addWidget(width_label, 0, 1)
        settings_layout.addWidget(self.image_width, 0, 0)
        settings_layout.addWidget(height_label, 1, 1)
        settings_layout.addWidget(self.image_height, 1, 0)
        settings_layout.addWidget(x_offset_label, 2, 1)
        settings_layout.addWidget(self.image_x_offset, 2, 0)
        settings_layout.addWidget(y_offset_label, 3, 1)
        settings_layout.addWidget(self.image_y_offset, 3, 0)
        settings_layout.addWidget(self.image_one_per_page, 4, 0, 1, 2)
        settings_layout.addWidget(self.image_center, 5, 0, 1, 2)

        settings_group.setLayout(settings_layout)

        # گروه خروجی
        output_group = QGroupBox("تنظیمات خروجی")
        output_group.setAlignment(Qt.AlignLeft)
        output_group.setStyleSheet(group_style)
        output_layout = QHBoxLayout()
        output_layout.setSpacing(10)

        output_label = QLabel("فایل خروجی:")
        output_label.setStyleSheet(label_style)
        output_label.setToolTip("مسیر ذخیره فایل PDF خروجی")

        self.image_output = QLineEdit()
        self.image_output.setPlaceholderText("مسیر ذخیره فایل PDF خروجی (مثال: output.pdf)")
        self.image_output.setAlignment(Qt.AlignLeft)
        self.image_output.setStyleSheet(input_style)
        self.image_output.setToolTip("مسیر فایل PDF خروجی را وارد یا انتخاب کنید")

        browse_output = QPushButton("انتخاب مسیر")
        browse_output.setStyleSheet(button_style)
        browse_output.setToolTip("انتخاب مسیر ذخیره فایل خروجی")
        browse_output.clicked.connect(lambda: self.image_output.setText(QFileDialog.getSaveFileName(self, "ذخیره فایل PDF", "", "فایل‌های PDF (*.pdf)")[0]))

        output_layout.addWidget(output_label)
        output_layout.addWidget(self.image_output)
        output_layout.addWidget(browse_output)
        output_group.setLayout(output_layout)

        # دکمه تبدیل
        convert_btn = QPushButton("تبدیل به PDF")
        convert_btn.setStyleSheet(convert_button_style)
        convert_btn.setToolTip("شروع فرآیند تبدیل تصاویر به PDF")
        convert_btn.clicked.connect(self.image_to_pdf)

        # افزودن گروه‌ها به چیدمان اصلی
        main_layout.addWidget(images_group)
        main_layout.addWidget(settings_group)
        main_layout.addWidget(output_group)
        main_layout.addWidget(convert_btn, alignment=Qt.AlignCenter)
        main_layout.addStretch()

        widget.setLayout(main_layout)

        # تنظیم فونت فارسی
        if hasattr(self, 'set_persian_font'):
            self.set_persian_font(widget)

        return widget

    def image_to_pdf(self):
        images = [self.image_list.item(i).text() for i in range(self.image_list.count())]
        output_file = self.image_output.text()
        if not images or not output_file:
            msg = QMessageBox()
            msg.setWindowTitle("هشدار")
            msg.setText("لطفاً تصاویر را اضافه کرده و فایل خروجی را مشخص کنید.")
            msg.setIcon(QMessageBox.Warning)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
            return
        try:
            c = canvas.Canvas(output_file, pagesize=letter)
            width, height = letter
            for i, img_path in enumerate(images):
                if i > 0 and self.image_one_per_page.isChecked():
                    c.showPage()
                img_w = self.image_width.value() or None
                img_h = self.image_height.value() or None
                x = (width - (img_w or width)) / 2 if self.image_center.isChecked() else self.image_x_offset.value()
                y = (height - (img_h or height)) / 2 if self.image_center.isChecked() else self.image_y_offset.value()
                c.drawImage(img_path, x, y, img_w, img_h)
            c.save()
            msg = QMessageBox()
            msg.setWindowTitle("موفقیت")
            msg.setText("تصاویر با موفقیت به PDF تبدیل شدند.")
            msg.setIcon(QMessageBox.Information)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
        except Exception as e:
            msg = QMessageBox()
            msg.setWindowTitle("خطا")
            msg.setText(f"خطا در تبدیل تصاویر به PDF: {str(e)}")
            msg.setIcon(QMessageBox.Critical)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()




    def create_reorder_pages_widget(self):
        # ایجاد ویجت اصلی
        widget = QWidget()
        widget.setLayoutDirection(Qt.RightToLeft)  # تنظیم جهت به چپ‌به‌راست
        main_layout = QVBoxLayout()
        main_layout.setAlignment(Qt.AlignLeft)
        main_layout.setSpacing(20)
        main_layout.setContentsMargins(20, 20, 20, 20)

        # استایل‌های مشترک
        input_style = """
            QLineEdit {
                padding: 10px;
                border-radius: 8px;
                border: 2px solid #e0e0e0;
                background-color: #ffffff;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QLineEdit:hover {
                border: 2px solid #007BFF;
                background-color: #f8faff;
            }
            QLineEdit:focus {
                border: 2px solid #0056b3;
                background-color: #ffffff;
            }
        """
        button_style = """
            QPushButton {
                padding: 10px;
                border-radius: 8px;
                background-color: #007BFF;
                color: white;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                font-weight: bold;
                border: none;
                min-width: 120px;
            }
            QPushButton:hover {
                background-color: #0056b3;
                border: 1px solid #004085;
            }
            QPushButton:pressed {
                background-color: #003d80;
            }
        """
        label_style = """
            QLabel {
                font-family: 'B Nazanin', 'Arial';
                font-size: 16px;
                font-weight: bold;
                color: #333333;
            }
        """
        group_style = """
            QGroupBox {
                border: 2px solid #e0e0e0;
                border-radius: 10px;
                padding: 15px;
                background-color: #ffffff;
                margin-top: 10px;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top left;
                padding: 5px 10px;
                color: #333333;
                font-weight: bold;
            }
        """
        reorder_button_style = """
            QPushButton {
                padding: 12px;
                border-radius: 10px;
                background-color: #28A745;
                color: white;
                font-family: 'B Nazanin', 'Arial';
                font-size: 16px;
                font-weight: bold;
                border: none;
                min-width: 200px;
            }
            QPushButton:hover {
                background-color: #218838;
                border: 1px solid #1e7e34;
            }
            QPushButton:pressed {
                background-color: #1a6b2d;
            }
        """

        # تنظیم فونت و استایل کلی
        widget.setStyleSheet("""
            QWidget {
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                background-color: #f5f5f5;
            }
        """)

        # گروه ورودی فایل
        input_group = QGroupBox("فایل ورودی")
        input_group.setAlignment(Qt.AlignLeft)
        input_group.setStyleSheet(group_style)
        input_layout = QHBoxLayout()
        input_layout.setSpacing(10)

        input_label = QLabel("فایل PDF ورودی:")
        input_label.setAlignment(Qt.AlignLeft)
        input_label.setStyleSheet(label_style)
        input_label.setToolTip("فایل PDF را برای مرتب‌سازی انتخاب کنید")

        self.reorder_input = QLineEdit()
        self.reorder_input.setPlaceholderText("فایل PDF را انتخاب کنید...")
        self.reorder_input.setAlignment(Qt.AlignLeft)
        self.reorder_input.setStyleSheet(input_style)
        self.reorder_input.setToolTip("مسیر فایل PDF ورودی را وارد یا انتخاب کنید")

        browse_btn = QPushButton("انتخاب فایل")
        browse_btn.setStyleSheet(button_style)
        browse_btn.setToolTip("انتخاب فایل PDF")
        browse_btn.clicked.connect(lambda: self.reorder_input.setText(QFileDialog.getOpenFileName(self, "انتخاب PDF", "", "فایل‌های PDF (*.pdf)")[0]))

        input_layout.addWidget(input_label)
        input_layout.addWidget(self.reorder_input)
        input_layout.addWidget(browse_btn)
        input_group.setLayout(input_layout)

        # گروه تنظیمات ترتیب
        order_group = QGroupBox("تنظیمات ترتیب")
        order_group.setAlignment(Qt.AlignLeft)
        order_group.setStyleSheet(group_style)
        order_layout = QHBoxLayout()
        order_layout.setSpacing(10)

        order_label = QLabel("ترتیب جدید صفحات:")
        order_label.setAlignment(Qt.AlignLeft)
        order_label.setStyleSheet(label_style)
        order_label.setToolTip("ترتیب جدید صفحات را وارد کنید (مثال: 3و1و2و4 یا 4تا6و1)")

        self.reorder_order = QLineEdit()
        self.reorder_order.setPlaceholderText("مثال: 3و1و2و4 یا 4تا6و1")
        self.reorder_order.setAlignment(Qt.AlignLeft)
        self.reorder_order.setStyleSheet(input_style)
        self.reorder_order.setToolTip("شماره صفحات را به ترتیب دلخواه وارد کنید")

        order_layout.addWidget(order_label)
        order_layout.addWidget(self.reorder_order)
        order_group.setLayout(order_layout)

        # گروه خروجی
        output_group = QGroupBox("فایل خروجی")
        output_group.setAlignment(Qt.AlignLeft)
        output_group.setStyleSheet(group_style)
        output_layout = QHBoxLayout()
        output_layout.setSpacing(10)

        output_label = QLabel("فایل خروجی:")
        output_label.setAlignment(Qt.AlignLeft)
        output_label.setStyleSheet(label_style)
        output_label.setToolTip("مسیر ذخیره فایل PDF خروجی را مشخص کنید")

        self.reorder_output = QLineEdit()
        self.reorder_output.setPlaceholderText("مسیر فایل خروجی را انتخاب کنید...")
        self.reorder_output.setAlignment(Qt.AlignLeft)
        self.reorder_output.setStyleSheet(input_style)
        self.reorder_output.setToolTip("مسیر ذخیره فایل PDF مرتب‌شده را وارد یا انتخاب کنید")

        browse_output = QPushButton("انتخاب مسیر")
        browse_output.setStyleSheet(button_style)
        browse_output.setToolTip("انتخاب مسیر ذخیره فایل خروجی")
        browse_output.clicked.connect(lambda: self.reorder_output.setText(QFileDialog.getSaveFileName(self, "ذخیره به عنوان", "", "فایل‌های PDF (*.pdf)")[0]))

        output_layout.addWidget(output_label)
        output_layout.addWidget(self.reorder_output)
        output_layout.addWidget(browse_output)
        output_group.setLayout(output_layout)

        # دکمه مرتب‌سازی
        reorder_btn = QPushButton("مرتب‌سازی صفحات")
        reorder_btn.setStyleSheet(reorder_button_style)
        reorder_btn.setToolTip("شروع فرآیند مرتب‌سازی صفحات")
        reorder_btn.clicked.connect(self.reorder_pages_func)

        # افزودن گروه‌ها به چیدمان اصلی
        main_layout.addWidget(input_group)
        main_layout.addWidget(order_group)
        main_layout.addWidget(output_group)
        main_layout.addWidget(reorder_btn, alignment=Qt.AlignCenter)
        main_layout.addStretch()

        widget.setLayout(main_layout)
        return widget

    def reorder_pages_func(self):
        input_file = self.reorder_input.text()
        output_file = self.reorder_output.text()
        order_text = self.reorder_order.text()

        # بررسی پر بودن فیلدها
        if not all([input_file, output_file, order_text]):
            msg = QMessageBox()
            msg.setWindowTitle("هشدار")
            msg.setText("لطفاً تمامی فیلدها را پر کنید.")
            msg.setIcon(QMessageBox.Warning)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
            return

        try:
            with open(input_file, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                total_pages = len(reader.pages)

                # تجزیه ترتیب صفحات
                new_order = []
                # جایگزینی فاصله‌ها و کاماها برای یکسان‌سازی
                order_text = order_text.replace(' ', 'و').replace(',', 'و')
                parts = order_text.split('و')

                for part in parts:
                    if 'تا' in part:
                        # پردازش محدوده‌ها (مثل 4تا6 یا 6تا4)
                        start, end = map(int, part.split('تا'))
                        if start > end:
                            # محدوده نزولی
                            new_order.extend(range(start, end - 1, -1))
                        else:
                            # محدوده صعودی
                            new_order.extend(range(start, end + 1))
                    else:
                        # پردازش اعداد منفرد
                        new_order.append(int(part))

                # اعتبارسنجی ترتیب جدید
                if len(new_order) != total_pages or set(new_order) != set(range(1, total_pages + 1)):
                    raise ValueError("ترتیب صفحات نامعتبر است. باید شامل همه صفحات باشد و تکراری نباشد.")

                # بازچینی صفحات
                writer = PyPDF2.PdfWriter()
                for page_num in new_order:
                    writer.add_page(reader.pages[page_num - 1])

                with open(output_file, 'wb') as out:
                    writer.write(out)

                # پیام موفقیت
                msg = QMessageBox()
                msg.setWindowTitle("موفقیت")
                msg.setText("صفحات با موفقیت مرتب شدند.")
                msg.setIcon(QMessageBox.Information)
                msg.setLayoutDirection(Qt.RightToLeft)
                msg.exec_()

        except FileNotFoundError:
            msg = QMessageBox()
            msg.setWindowTitle("خطا")
            msg.setText("فایل ورودی یافت نشد.")
            msg.setIcon(QMessageBox.Critical)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
        except ValueError as ve:
            msg = QMessageBox()
            msg.setWindowTitle("خطا")
            msg.setText(f"خطا در ترتیب صفحات: {str(ve)}")
            msg.setIcon(QMessageBox.Critical)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
        except Exception as e:
            msg = QMessageBox()
            msg.setWindowTitle("خطا")
            msg.setText(f"خطا در مرتب‌سازی صفحات: {str(e)}")
            msg.setIcon(QMessageBox.Critical)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()





    def create_scan_pdf_widget(self):
        # ایجاد ویجت اصلی
        widget = QWidget()
        widget.setLayoutDirection(Qt.RightToLeft)  # تنظیم جهت به چپ‌به‌راست
        main_layout = QVBoxLayout()
        main_layout.setAlignment(Qt.AlignLeft)
        main_layout.setSpacing(20)
        main_layout.setContentsMargins(30, 30, 30, 30)  # حاشیه‌های مدرن

        # استایل‌های مشترک
        input_style = """
            QLineEdit {
                padding: 10px;
                border-radius: 8px;
                border: 2px solid #e0e0e0;
                background-color: #ffffff;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QLineEdit:hover {
                border: 2px solid #007BFF;
                background-color: #f8faff;
            }
            QLineEdit:focus {
                border: 2px solid #0056b3;
                background-color: #ffffff;
            }
        """
        button_style = """
            QPushButton {
                padding: 10px;
                border-radius: 8px;
                background-color: #007BFF;
                color: white;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                font-weight: bold;
                border: none;
                min-width: 120px;
            }
            QPushButton:hover {
                background-color: #0056b3;
                border: 1px solid #004085;
            }
            QPushButton:pressed {
                background-color: #003d80;
            }
        """
        scan_button_style = """
            QPushButton {
                padding: 12px;
                border-radius: 10px;
                background-color: #28A745;
                color: white;
                font-family: 'B Nazanin', 'Arial';
                font-size: 16px;
                font-weight: bold;
                border: none;
                min-width: 200px;
            }
            QPushButton:hover {
                background-color: #218838;
                border: 1px solid #1e7e34;
            }
            QPushButton:pressed {
                background-color: #1a6b2d;
            }
        """
        cancel_button_style = """
            QPushButton {
                padding: 10px;
                border-radius: 8px;
                background-color: #DC3545;
                color: white;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                font-weight: bold;
                border: none;
                min-width: 120px;
            }
            QPushButton:hover {
                background-color: #C82333;
                border: 1px solid #BD2130;
            }
            QPushButton:pressed {
                background-color: #B21F2D;
            }
        """
        label_style = """
            QLabel {
                font-family: 'B Nazanin', 'Arial';
                font-size: 16px;
                font-weight: bold;
                color: #333333;
            }
        """
        group_style = """
            QGroupBox {
                border: 2px solid #e0e0e0;
                border-radius: 10px;
                padding: 15px;
                background-color: #ffffff;
                margin-top: 10px;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top left;
                padding: 5px 10px;
                color: #333333;
                font-weight: bold;
            }
        """
        combo_style = """
            QComboBox {
                padding: 10px;
                border-radius: 8px;
                border: 2px solid #e0e0e0;
                background-color: #ffffff;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QComboBox:hover {
                border: 2px solid #007BFF;
                background-color: #f8faff;
            }
        """
        spinbox_style = """
            QSpinBox {
                padding: 10px;
                border-radius: 8px;
                border: 2px solid #e0e0e0;
                background-color: #ffffff;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QSpinBox:hover {
                border: 2px solid #007BFF;
                background-color: #f8faff;
            }
        """
        progress_style = """
            QProgressBar {
                border: 2px solid #e0e0e0;
                border-radius: 8px;
                text-align: center;
                background-color: #f0f0f0;
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
            }
            QProgressBar::chunk {
                background-color: #28A745;
                border-radius: 6px;
            }
        """

        # تنظیم فونت و استایل کلی
        widget.setStyleSheet("""
            QWidget {
                font-family: 'B Nazanin', 'Arial';
                font-size: 14px;
                background-color: #f5f5f5;
            }
        """)

        # گروه ورودی فایل
        input_group = QGroupBox("فایل PDF ورودی")
        input_group.setAlignment(Qt.AlignLeft)
        input_group.setStyleSheet(group_style)
        input_layout = QHBoxLayout()
        input_layout.setSpacing(10)

        input_label = QLabel("فایل PDF ورودی:")
        input_label.setStyleSheet(label_style)
        input_label.setAlignment(Qt.AlignLeft)
        input_label.setToolTip("فایل PDF را برای اسکن انتخاب کنید")

        self.scan_input = QLineEdit()
        self.scan_input.setPlaceholderText("فایل PDF را انتخاب کنید...")
        self.scan_input.setAlignment(Qt.AlignLeft)
        self.scan_input.setStyleSheet(input_style)
        self.scan_input.setToolTip("مسیر فایل PDF ورودی را وارد یا انتخاب کنید")

        browse_btn = QPushButton("انتخاب فایل")
        browse_btn.setStyleSheet(button_style)
        browse_btn.setToolTip("انتخاب فایل PDF")
        browse_btn.clicked.connect(lambda: self.scan_input.setText(QFileDialog.getOpenFileName(self, "انتخاب PDF", "", "فایل‌های PDF (*.pdf)")[0]))

        input_layout.addWidget(input_label)
        input_layout.addWidget(self.scan_input)
        input_layout.addWidget(browse_btn)
        input_group.setLayout(input_layout)

        # گروه تنظیمات OCR
        ocr_settings_group = QGroupBox("تنظیمات OCR")
        ocr_settings_group.setAlignment(Qt.AlignLeft)
        ocr_settings_group.setStyleSheet(group_style)
        ocr_settings_layout = QGridLayout()
        ocr_settings_layout.setSpacing(10)

        language_label = QLabel("زبان OCR:")
        language_label.setStyleSheet(label_style)
        language_label.setAlignment(Qt.AlignLeft)
        language_label.setToolTip("زبان مورد نظر برای اسکن متن")

        self.ocr_language = QComboBox()
        self.ocr_language.addItems(["فارسی", "انگلیسی", "فارسی و انگلیسی"])
        self.ocr_language.setStyleSheet(combo_style)
        self.ocr_language.setToolTip("زبان OCR را انتخاب کنید")

        dpi_label = QLabel("DPI اسکن:")
        dpi_label.setStyleSheet(label_style)
        dpi_label.setAlignment(Qt.AlignLeft)
        dpi_label.setToolTip("وضوح تصویر برای اسکن")

        self.ocr_dpi = QSpinBox()
        self.ocr_dpi.setRange(100, 600)
        self.ocr_dpi.setValue(500)
        self.ocr_dpi.setStyleSheet(spinbox_style)
        self.ocr_dpi.setToolTip("مقدار DPI را تنظیم کنید (100 تا 600)")

        ocr_settings_layout.addWidget(language_label, 0, 0)
        ocr_settings_layout.addWidget(self.ocr_language, 0, 1)
        ocr_settings_layout.addWidget(dpi_label, 1, 0)
        ocr_settings_layout.addWidget(self.ocr_dpi, 1, 1)
        ocr_settings_group.setLayout(ocr_settings_layout)

        # گروه خروجی
        output_group = QGroupBox("فایل خروجی")
        output_group.setAlignment(Qt.AlignLeft)
        output_group.setStyleSheet(group_style)
        output_layout = QHBoxLayout()
        output_layout.setSpacing(10)

        output_label = QLabel("فایل Word خروجی:")
        output_label.setStyleSheet(label_style)
        output_label.setAlignment(Qt.AlignLeft)
        output_label.setToolTip("محل ذخیره فایل Word خروجی")

        self.scan_output = QLineEdit()
        self.scan_output.setPlaceholderText("مسیر فایل Word خروجی را انتخاب کنید...")
        self.scan_output.setAlignment(Qt.AlignLeft)
        self.scan_output.setStyleSheet(input_style)
        self.scan_output.setToolTip("مسیر ذخیره فایل Word خروجی را وارد یا انتخاب کنید")

        browse_output = QPushButton("انتخاب مسیر")
        browse_output.setStyleSheet(button_style)
        browse_output.setToolTip("انتخاب مسیر ذخیره فایل خروجی")
        browse_output.clicked.connect(lambda: self.scan_output.setText(QFileDialog.getSaveFileName(self, "ذخیره به عنوان", "", "فایل‌های Word (*.docx)")[0]))

        output_layout.addWidget(output_label)
        output_layout.addWidget(self.scan_output)
        output_layout.addWidget(browse_output)
        output_group.setLayout(output_layout)

        # نوار پیشرفت
        progress_group = QGroupBox("پیشرفت")
        progress_group.setAlignment(Qt.AlignLeft)
        progress_group.setStyleSheet(group_style)
        progress_layout = QVBoxLayout()

        self.scan_progress = QProgressBar()
        self.scan_progress.setValue(0)
        self.scan_progress.setStyleSheet(progress_style)
        self.scan_progress.setToolTip("پیشرفت فرآیند اسکن و استخراج متن")

        progress_layout.addWidget(self.scan_progress)
        progress_group.setLayout(progress_layout)

        # دکمه‌های عملیاتی
        button_layout = QHBoxLayout()
        button_layout.setSpacing(10)

        scan_btn = QPushButton("اسکن و استخراج متن")
        scan_btn.setStyleSheet(scan_button_style)
        scan_btn.setToolTip("شروع فرآیند اسکن و استخراج متن")
        scan_btn.clicked.connect(self.start_scan)

        self.cancel_btn = QPushButton("انصراف")
        self.cancel_btn.setStyleSheet(cancel_button_style)
        self.cancel_btn.setToolTip("لغو فرآیند اسکن")
        self.cancel_btn.clicked.connect(self.cancel_scan)
        self.cancel_btn.setEnabled(False)

        button_layout.addWidget(scan_btn)
        button_layout.addWidget(self.cancel_btn)

        # افزودن گروه‌ها به چیدمان اصلی
        main_layout.addWidget(input_group)
        main_layout.addWidget(ocr_settings_group)
        main_layout.addWidget(output_group)
        main_layout.addWidget(progress_group)
        main_layout.addLayout(button_layout)
        main_layout.addStretch()

        widget.setLayout(main_layout)
        return widget

    def find_tesseract(self):
        possible_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]
        for path in possible_paths:
            if os.path.exists(path):
                return path
        try:
            result = subprocess.run(["tesseract", "--version"], capture_output=True, text=True, check=True)
            if result.returncode == 0:
                return "tesseract"
        except (subprocess.CalledProcessError, FileNotFoundError):
            pass
        file_path, _ = QFileDialog.getOpenFileName(None, "انتخاب فایل اجرایی tesseract", "",
                                                "فایل‌های اجرایی (*.exe);;همه فایل‌ها (*)")
        return file_path if file_path else None

    def check_tesseract(self):
        tesseract_path = self.find_tesseract()
        if tesseract_path and tesseract_path != "tesseract":
            tesseract_dir = os.path.dirname(tesseract_path)
            current_path = os.environ.get('PATH', '')
            if tesseract_dir not in current_path.split(os.pathsep):
                try:
                    new_path = f"{current_path}{os.pathsep}{tesseract_dir}"
                    subprocess.run(f'setx PATH "{new_path}"', check=True, shell=True)
                    os.environ['PATH'] = new_path
                except Exception as e:
                    QMessageBox.critical(self, "خطا", f"خطا در افزودن tesseract به PATH: {e}")
                    return False
        if not tesseract_path:
            self.show_tesseract_install_dialog()
            return False
        return True

    def show_tesseract_install_dialog(self):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Warning)
        msg.setText("Tesseract یافت نشد.")
        msg.setInformativeText("برای اسکن PDF، لطفاً Tesseract OCR را نصب کنید.")
        msg.setWindowTitle("نیاز به Tesseract")
        msg.setLayoutDirection(Qt.RightToLeft)
        button_style = """
            QPushButton {
                padding: 8px; 
                border-radius: 5px; 
                background-color: #007BFF; 
                color: white;
            }
            QPushButton:hover {
                background-color: #0056b3;
            }
        """
        ok_btn = msg.addButton(QMessageBox.Ok)
        ok_btn.setStyleSheet(button_style)
        msg.exec_()

    def start_scan(self):
        if not self.check_tesseract():
            return
        input_file = self.scan_input.text().strip()
        output_file = self.scan_output.text().strip()
        language = self.ocr_language.currentText()
        dpi = self.ocr_dpi.value()
        if not all([input_file, output_file]):
            msg = QMessageBox()
            msg.setWindowTitle("هشدار")
            msg.setText("لطفاً فیلدهای ورودی و خروجی را پر کنید.")
            msg.setIcon(QMessageBox.Warning)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
            return
        if not input_file.endswith('.pdf'):
            msg = QMessageBox()
            msg.setWindowTitle("هشدار")
            msg.setText("لطفاً یک فایل PDF معتبر انتخاب کنید.")
            msg.setIcon(QMessageBox.Warning)
            msg.setLayoutDirection(Qt.RightToLeft)
            msg.exec_()
            return
        if not output_file.endswith('.docx'):
            output_file += '.docx'
        scan_btn = self.findChild(QPushButton, "اسکن و استخراج متن")
        if scan_btn:
            scan_btn.setEnabled(False)
        self.cancel_btn.setEnabled(True)
        self.scan_thread = QThread()
        self.scan_worker = self.OCRWorker(input_file, output_file, language, dpi)
        self.scan_worker.moveToThread(self.scan_thread)
        self.scan_thread.started.connect(self.scan_worker.run)
        self.scan_worker.progress.connect(self.update_progress)
        self.scan_worker.finished.connect(self.scan_finished)
        self.scan_worker.error.connect(self.scan_error)
        self.scan_thread.start()

    def update_progress(self, value):
        self.scan_progress.setValue(value)

    def scan_finished(self):
        self.scan_progress.setValue(100)
        scan_btn = self.findChild(QPushButton, "اسکن و استخراج متن")
        if scan_btn:
            scan_btn.setEnabled(True)
        self.cancel_btn.setEnabled(False)
        self.scan_progress.setValue(0)
        msg = QMessageBox()
        msg.setWindowTitle("موفقیت")
        msg.setText("متن با موفقیت استخراج و به سند Word تبدیل شد.")
        msg.setIcon(QMessageBox.Information)
        msg.setLayoutDirection(Qt.RightToLeft)
        msg.exec_()
        self.scan_thread.quit()
        self.scan_thread.wait()

    def scan_error(self, error_msg):
        scan_btn = self.findChild(QPushButton, "اسکن و استخراج متن")
        if scan_btn:
            scan_btn.setEnabled(True)
        self.cancel_btn.setEnabled(False)
        self.scan_progress.setValue(0)
        msg = QMessageBox()
        msg.setWindowTitle("خطا")
        msg.setText(f"خطا در اسکن و استخراج متن: {error_msg}")
        msg.setIcon(QMessageBox.Critical)
        msg.setLayoutDirection(Qt.RightToLeft)
        msg.exec_()
        self.scan_thread.quit()
        self.scan_thread.wait()

    def cancel_scan(self):
        if hasattr(self, 'scan_worker'):
            self.scan_worker.stop()
        scan_btn = self.findChild(QPushButton, "اسکن و استخراج متن")
        if scan_btn:
            scan_btn.setEnabled(True)
        self.cancel_btn.setEnabled(False)
        self.scan_progress.setValue(0)
        msg = QMessageBox()
        msg.setWindowTitle("انصراف")
        msg.setText("عملیات اسکن و استخراج متن لغو شد.")
        msg.setIcon(QMessageBox.Information)
        msg.setLayoutDirection(Qt.RightToLeft)
        msg.exec_()
        if hasattr(self, 'scan_thread'):
            self.scan_thread.quit()
            self.scan_thread.wait()

    class OCRWorker(QObject):
        progress = pyqtSignal(int)
        finished = pyqtSignal()
        error = pyqtSignal(str)

        def __init__(self, input_file, output_file, language, dpi):
            super().__init__()
            self.input_file = input_file
            self.output_file = output_file
            self.language = language
            self.dpi = dpi
            self._is_running = True

        def run(self):
            try:
                if self.language == "فارسی":
                    ocr_lang = "fas"
                elif self.language == "انگلیسی":
                    ocr_lang = "eng"
                else:
                    ocr_lang = "fas+eng"
                doc = fitz.open(self.input_file)
                word_doc = Document()
                total_pages = len(doc)
                for page_num in range(total_pages):
                    if not self._is_running:
                        self.error.emit("عملیات توسط کاربر لغو شد.")
                        return
                    page = doc.load_page(page_num)
                    pix = page.get_pixmap(dpi=self.dpi)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    img_cv = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
                    gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
                    _, binarized = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                    denoised = cv2.fastNlMeansDenoising(binarized, None, 10, 7, 21)
                    img_processed = Image.fromarray(denoised)
                    custom_config = r'--psm 6 --oem 3 -c preserve_interword_spaces=1'
                    text = image_to_string(img_processed, lang=ocr_lang, config=custom_config)
                    paragraph = word_doc.add_paragraph()
                    run = paragraph.add_run(f"صفحه {page_num + 1}:\n{text}\n\n")
                    run.font.name = 'B Nazanin'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                    run.font.size = Pt(12)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.Left
                    self.progress.emit(int((page_num + 1) / total_pages * 100))
                if self._is_running:
                    word_doc.save(self.output_file)
                    self.finished.emit()
            except Exception as e:
                self.error.emit(str(e))

        def stop(self):
            self._is_running = False



